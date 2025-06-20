import os
import sys
import pandas as pd
import numpy as np
from PyQt5.QtWidgets import (QApplication, QMainWindow, QFileDialog, QTabWidget, 
                            QWidget, QVBoxLayout, QPushButton, QLabel, QTableView,
                            QHBoxLayout, QMessageBox, QProgressBar, QComboBox,
                            QListWidget, QGridLayout, QSplitter, QSizePolicy, QDialog, QDialogButtonBox, QAbstractItemView,
                            QSpinBox, QFrame, QTableWidget, QTableWidgetItem, QScrollArea)
from PyQt5.QtCore import Qt, QSize, QThread, pyqtSignal, QTimer
from PyQt5 import QtGui
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.backends.backend_qt5agg import FigureCanvas
from matplotlib.figure import Figure
import time

# Import XML processing functions from existing code
from xml_processor import extract_beyanname_fixed, process_all_xml_files, process_multiple_xml_files, merge_dataframes
from custom_widgets import PandasModel, DataFrameViewer, CheckResultsWidget

# Import analysis modules
from analysis_modules import (
    # Basic checks
    calculate_basic_stats, check_missing_values, check_duplicate_rows,
    check_value_consistency, check_numeric_range,
    
    # Analysis functions
    check_weight_consistency, check_currency_values, check_vergi_consistency,
    check_alici_satici_relationship,
    check_gtip_urun_kodu_consistency, check_rarely_used_currency,
    check_rarely_used_origin_country, check_rarely_used_origin_country_by_sender_gtip,
    check_rarely_used_payment_method,
    check_unit_price_increase, check_kdv_consistency,
    check_domestic_expense_variation, check_foreign_expense_variation,
    check_supalan_depolama_kontrol,
    check_tedarikci_beyan_kontrol,
    
    # Summary functions
    create_gtip_summary, create_country_summary, create_rejim_summary,
    create_gtip_country_cross,
    
    # Chart functions
    create_bar_chart, create_pie_chart
)

# Import specific modules
from analysis_modules.islem_niteligi_tutarlilik import kontrol_islem_niteligi_tutarlilik
from analysis_modules.kkdf_kontrol import check_kkdf_kontrol
from analysis_modules.gozetim_kontrol import check_gozetim_kontrol

# Örnekleme modülünü içe aktar
from sampling import BeyannameSampling

# QThread sınıfı ekliyorum (Excel işlemini arka planda yapacak)
class ExcelExportThread(QThread):
    # Sinyaller
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(bool, str, str)  # başarı/başarısız, mesaj, dosya yolu
    
    def __init__(self, sampling_tool, file_path):
        super().__init__()
        self.sampling_tool = sampling_tool
        self.file_path = file_path
        self.is_cancelled = False
    
    def run(self):
        try:
            # İlerleme bilgisi
            self.progress.emit(20, "Excel dosyası hazırlanıyor...")
            
            # Güvenlik kontrolü - timeout ekleme
            self.timer = QTimer()
            self.timer.setSingleShot(True)
            self.timer.timeout.connect(self.handle_timeout)
            self.timer.start(60000)  # 60 saniyelik timeout
            
            # Excel'e aktarma
            try:
                output_path = self.sampling_tool.export_to_excel(self.file_path)
                if self.is_cancelled:
                    return
                
                self.progress.emit(70, "Excel formatlanıyor...")
                
                # Excel dosyasını formatla
                self.sampling_tool.format_excel_report(output_path)
                if self.is_cancelled:
                    return
                
                # İşlem başarılı
                self.finished.emit(True, f"Örnekleme sonuçları Excel'e aktarıldı: {output_path}", output_path)
                
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"Excel aktarma detaylı hata: {error_details}")
                
                error_msg = str(e)
                # Hata mesajını daha anlaşılır hale getir
                if "PermissionError" in error_details:
                    error_msg = f"Dosya erişim hatası: {self.file_path}\nDosya başka bir program tarafından kullanılıyor olabilir."
                elif "FileNotFoundError" in error_details:
                    error_msg = f"Dosya yolu bulunamadı: {self.file_path}"
                
                self.finished.emit(False, f"Excel'e aktarma hatası: {error_msg}", "")
            
            # Timer'ı durdur
            self.timer.stop()
                
        except Exception as e:
            # Genel hata durumunda
            import traceback
            print(f"Excel thread beklenmeyen hata: {str(e)}")
            print(traceback.format_exc())
            self.finished.emit(False, f"Beklenmeyen hata: {str(e)}", "")
    
    def handle_timeout(self):
        """İşlem zaman aşımına uğradığında çağrılır"""
        self.is_cancelled = True
        self.terminate()  # Thread'i sonlandır
        self.finished.emit(False, "İşlem zaman aşımına uğradı. Excel çok büyük olabilir.", "")
    
    def cancel(self):
        """İşlemi iptal et"""
        self.is_cancelled = True
        self.timer.stop()
        self.terminate()

class CustomsCheckApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Beyanname Kontrol Uygulaması")
        self.setGeometry(100, 100, 1200, 800)
        
        # Store loaded data
        self.all_dataframes = {}
        self.current_df = None
        self.merged_df = None  # Birleştirilmiş tüm veriler için
        
        # Excel thread değişkeni
        self.excel_thread = None
        self.cancel_dialog = None
        
        # Set application style
        self.apply_modern_style()
        
        self.init_ui()
        
        # Stabilite ayarları
        self.setup_application_stability()
        
        # Kısayol tuşlarını ayarla
        self.setup_shortcuts()
    
    def apply_modern_style(self):
        """Apply modern Apple-style theme to the application"""
        style_sheet = """
        /* Modern Apple-like Theme */
        QMainWindow {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #f8fafb, stop:1 #e8edf2);
            color: #1c1e21;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
        }
        
        /* Modern Tab Widget */
        QTabWidget {
            background: transparent;
            border: none;
        }
        
        QTabWidget::pane {
            background: rgba(255, 255, 255, 0.95);
            border: 1px solid rgba(0, 0, 0, 0.1);
            border-radius: 12px;
            margin: 0px;
            padding: 0px;
        }
        
        QTabWidget::tab-bar {
            left: 15px;
            top: 5px;
        }
        
        QTabBar::tab {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(255, 255, 255, 0.9), stop:1 rgba(248, 250, 251, 0.8));
            border: 1px solid rgba(0, 0, 0, 0.08);
            border-bottom: none;
            border-radius: 8px 8px 0px 0px;
            padding: 12px 24px;
            margin-right: 2px;
            margin-top: 3px;
            color: #4a5568;
            font-weight: 500;
            font-size: 13px;
            min-width: 120px;
        }
        
        QTabBar::tab:selected {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f7fafc);
            color: #2d3748;
            font-weight: 600;
            border-color: rgba(0, 0, 0, 0.1);
            margin-top: 0px;
        }
        
        QTabBar::tab:hover:!selected {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(255, 255, 255, 0.95), stop:1 rgba(247, 250, 252, 0.9));
            border-color: rgba(0, 0, 0, 0.12);
        }
        
        /* Modern Buttons */
        QPushButton {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4299e1, stop:1 #3182ce);
            border: 1px solid rgba(49, 130, 206, 0.6);
            border-radius: 8px;
            padding: 10px 20px;
            color: white;
            font-weight: 600;
            font-size: 13px;
            min-height: 16px;
        }
        
        QPushButton:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3182ce, stop:1 #2c5aa0);
            border-color: rgba(44, 90, 160, 0.8);
        }
        
        QPushButton:pressed {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #2c5aa0, stop:1 #2a4d8f);
        }
        
        QPushButton[class="primary"] {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4299e1, stop:1 #3182ce);
            border-color: rgba(49, 130, 206, 0.6);
        }
        
        QPushButton[class="success"] {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #48bb78, stop:1 #38a169);
            border-color: rgba(56, 161, 105, 0.6);
        }
        
        QPushButton[class="danger"] {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #f56565, stop:1 #e53e3e);
            border-color: rgba(229, 62, 62, 0.6);
        }
        
        /* Modern Inputs */
        QComboBox, QSpinBox, QLineEdit {
            background: rgba(255, 255, 255, 0.9);
            border: 1px solid rgba(226, 232, 240, 0.8);
            border-radius: 6px;
            padding: 8px 12px;
            font-size: 13px;
            color: #2d3748;
        }
        
        QComboBox:focus, QSpinBox:focus, QLineEdit:focus {
            border-color: #4299e1;
            background: rgba(255, 255, 255, 0.95);
        }
        
        /* Modern Tables */
        QTableView, QTableWidget {
            background: rgba(255, 255, 255, 0.95);
            border: 1px solid rgba(226, 232, 240, 0.8);
            border-radius: 8px;
            gridline-color: rgba(226, 232, 240, 0.6);
            selection-background-color: rgba(66, 153, 225, 0.2);
        }
        
        QTableView::item {
            padding: 8px;
            border: none;
        }
        
        QHeaderView::section {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(247, 250, 252, 0.95), stop:1 rgba(237, 242, 247, 0.9));
            border: 1px solid rgba(226, 232, 240, 0.8);
            border-radius: 0px;
            padding: 10px 8px;
            font-weight: 600;
            font-size: 12px;
            color: #4a5568;
        }
        
        /* Modern Progress Bar */
        QProgressBar {
            background: rgba(237, 242, 247, 0.8);
            border: 1px solid rgba(203, 213, 224, 0.6);
            border-radius: 8px;
            text-align: center;
            font-weight: 500;
            font-size: 12px;
            color: #2d3748;
        }
        
        QProgressBar::chunk {
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #4299e1, stop:1 #63b3ed);
            border-radius: 6px;
        }
        
        /* Modern Labels */
        QLabel {
            color: #2d3748;
            font-size: 13px;
        }
        
        /* Modern Scroll Bars */
        QScrollBar:vertical {
            background: rgba(237, 242, 247, 0.6);
            width: 12px;
            border-radius: 6px;
            margin: 0px;
        }
        
        QScrollBar::handle:vertical {
            background: rgba(160, 174, 192, 0.8);
            border-radius: 6px;
            min-height: 20px;
        }
        
        QScrollBar::handle:vertical:hover {
            background: rgba(113, 128, 150, 0.9);
        }
        
        /* Modern Group Box */
        QGroupBox {
            background: rgba(255, 255, 255, 0.9);
            border: 1px solid rgba(226, 232, 240, 0.8);
            border-radius: 12px;
            margin-top: 20px;
            padding-top: 10px;
            font-weight: 600;
            font-size: 14px;
            color: #2d3748;
        }
        
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 15px;
            top: 8px;
            background: rgba(255, 255, 255, 0.95);
            padding: 4px 12px;
            border-radius: 6px;
            border: 1px solid rgba(226, 232, 240, 0.8);
        }
        
        /* Modern Menu Bar */
        QMenuBar {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(255, 255, 255, 0.98), stop:1 rgba(248, 250, 251, 0.95));
            border-bottom: 1px solid rgba(226, 232, 240, 0.8);
            padding: 4px 8px;
            font-size: 13px;
            font-weight: 500;
        }
        
        QMenuBar::item {
            background: transparent;
            padding: 8px 16px;
            border-radius: 6px;
            color: #4a5568;
        }
        
        QMenuBar::item:selected {
            background: rgba(190, 227, 248, 0.6);
            color: #2d3748;
        }
        
        /* Modern Context Menu */
        QMenu {
            background: rgba(255, 255, 255, 0.98);
            border: 1px solid rgba(0, 0, 0, 0.1);
            border-radius: 8px;
            padding: 4px;
            font-size: 13px;
        }
        
        QMenu::item {
            background: transparent;
            padding: 8px 16px;
            border-radius: 4px;
            color: #2d3748;
        }
        
        QMenu::item:selected {
            background: rgba(190, 227, 248, 0.8);
        }
        
        QMenu::separator {
            height: 1px;
            background: rgba(226, 232, 240, 0.8);
            margin: 4px 8px;
        }
        """
        self.setStyleSheet(style_sheet)
        
    def init_ui(self):
        """Initialize the user interface"""
        # Main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # Top control panel
        control_panel = QWidget()
        control_panel.setObjectName("controlPanel")
        control_panel.setStyleSheet("""
            #controlPanel {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(255, 255, 255, 0.98), stop:1 rgba(248, 250, 251, 0.95));
                border: 1px solid rgba(226, 232, 240, 0.8);
                border-radius: 16px;
                padding: 15px;
                margin: 10px;
            }
        """)
        control_layout = QHBoxLayout(control_panel)
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(10)
        
        # Import XML button
        import_xml_btn = QPushButton("📁 XML Dosyası İçe Aktar")
        import_xml_btn.setProperty("class", "primary")
        import_xml_btn.clicked.connect(self.import_xml)
        control_layout.addWidget(import_xml_btn)
        
        # Import XML folder button
        import_folder_btn = QPushButton("📂 XML Klasörü İçe Aktar")
        import_folder_btn.setProperty("class", "primary")
        import_folder_btn.clicked.connect(self.import_xml_folder)
        control_layout.addWidget(import_folder_btn)
        
        # Batch import XML folder button
        batch_import_btn = QPushButton("⚡ Toplu XML İçe Aktar")
        batch_import_btn.setProperty("class", "success")
        batch_import_btn.clicked.connect(self.import_xml_folder)
        control_layout.addWidget(batch_import_btn)
        
        # Merge all dataframes button
        merge_all_btn = QPushButton("🔗 Tümünü Birleşik Göster")
        merge_all_btn.setProperty("class", "warning")
        merge_all_btn.clicked.connect(self.show_merged_dataframes)
        control_layout.addWidget(merge_all_btn)
        
        # File selector dropdown (will be populated after import)
        self.file_selector = QComboBox()
        self.file_selector.setMinimumWidth(300)
        self.file_selector.currentIndexChanged.connect(self.change_active_file)
        control_layout.addWidget(QLabel("Aktif Dosya:"))
        control_layout.addWidget(self.file_selector)
        
        # Status label
        self.status_label = QLabel("Henüz veri yüklenmedi")
        self.status_label.setStyleSheet("color: #757575; font-style: italic;")
        control_layout.addStretch()
        control_layout.addWidget(self.status_label)
        
        main_layout.addWidget(control_panel)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        # Tab widget for different views
        self.tabs = QTabWidget()
        self.tabs.setContentsMargins(0, 10, 0, 0)
        
        # Data tab
        self.data_tab = QWidget()
        data_layout = QVBoxLayout(self.data_tab)
        data_layout.setContentsMargins(10, 10, 10, 10)
        
        # Use DataFrameViewer instead of QTableView
        self.data_viewer = DataFrameViewer()
        data_layout.addWidget(self.data_viewer)
        
        self.tabs.addTab(self.data_tab, "Veri Görünümü")
        
        # Initialize sampling tool
        self.sampling_tool = BeyannameSampling()
        
        # Analysis tab (sol panel + sağ panel)
        self.analysis_tab = QWidget()
        analysis_layout = QHBoxLayout(self.analysis_tab)
        analysis_layout.setContentsMargins(0, 0, 0, 0)
        
        # Sol panel - Analiz butonları
        self.left_panel = QFrame()
        self.left_panel.setFixedWidth(250)
        left_layout = QVBoxLayout(self.left_panel)
        left_layout.setContentsMargins(10, 10, 10, 10)
        left_layout.setSpacing(8)
        
        # Tümü Çalıştır butonu (en üstte)
        btn_run_all = QPushButton("🚀 TÜMÜ ÇALIŞTIR")
        btn_run_all.setProperty("class", "primary")
        btn_run_all.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ff6b35, stop:1 #f7931e);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px;
                font-weight: bold;
                font-size: 14px;
                margin: 5px 0;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ff8c42, stop:1 #ffb347);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #e55a2b, stop:1 #e8831c);
            }
        """)
        btn_run_all.clicked.connect(self.run_all_analyses_and_export)
        left_layout.addWidget(btn_run_all)
        
        # Ayırıcı çizgi
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        separator.setStyleSheet("color: #ddd; margin: 10px 0;")
        left_layout.addWidget(separator)
        
        # Word Rapor butonu
        btn_word_report = QPushButton("📄 WORD RAPORU OLUŞTUR")
        btn_word_report.setProperty("class", "success")
        btn_word_report.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #28a745, stop:1 #20c997);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px;
                font-weight: bold;
                font-size: 14px;
                margin: 5px 0;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #34ce57, stop:1 #36d399);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #1e7e34, stop:1 #17a2b8);
            }
        """)
        btn_word_report.clicked.connect(self.create_word_report)
        left_layout.addWidget(btn_word_report)
        
        # Ayırıcı çizgi
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setFrameShadow(QFrame.Sunken)
        separator2.setStyleSheet("color: #ddd; margin: 10px 0;")
        left_layout.addWidget(separator2)
        
        # Temel kontroller başlığı
        basic_label = QLabel("Temel Kontroller")
        basic_label.setStyleSheet("font-weight: bold; color: #333; margin: 10px 0 5px 0;")
        left_layout.addWidget(basic_label)
        
        btn_missing = QPushButton("📊 Eksik Değerler")
        btn_missing.clicked.connect(self.check_missing_values)
        left_layout.addWidget(btn_missing)
        
        left_layout.addWidget(QLabel("Temel Kontroller:"))
        btn_islem = QPushButton("🔍 İşlem Niteliği")
        btn_islem.clicked.connect(self.check_islem_niteligi_consistency)
        left_layout.addWidget(btn_islem)
        
        btn_gtip_tanim_detail = QPushButton("🔍 GTİP-Tanım Detay")
        btn_gtip_tanim_detail.clicked.connect(self.check_gtip_tanim_detail)
        left_layout.addWidget(btn_gtip_tanim_detail)
        
        btn_gtip_urun = QPushButton("🏷️ GTİP-Ürün Kodu")
        btn_gtip_urun.clicked.connect(self.check_gtip_urun_kodu_consistency)
        left_layout.addWidget(btn_gtip_urun)
        
        btn_alici_satici = QPushButton("🤝 Alıcı-Satıcı İlişkisi")
        btn_alici_satici.clicked.connect(self.check_alici_satici_relationship)
        left_layout.addWidget(btn_alici_satici)
        
        left_layout.addSpacing(10)
        left_layout.addWidget(QLabel("İstatistiksel Kontroller:"))
        
        btn_nadir_doviz = QPushButton("💱 Nadir Döviz")
        btn_nadir_doviz.clicked.connect(self.check_rarely_used_currency)
        left_layout.addWidget(btn_nadir_doviz)
        
        btn_nadir_ulke = QPushButton("🌍 Nadir Menşe Ülke")
        btn_nadir_ulke.clicked.connect(self.check_rarely_used_origin_country)
        left_layout.addWidget(btn_nadir_ulke)
        
        btn_nadir_ulke_gtip = QPushButton("🔍 Gönderici-GTİP Nadir Menşe")
        btn_nadir_ulke_gtip.clicked.connect(self.check_rarely_used_origin_country_by_sender_gtip)
        left_layout.addWidget(btn_nadir_ulke_gtip)
        
        btn_nadir_odeme = QPushButton("💳 Nadir Ödeme Şekli")
        btn_nadir_odeme.clicked.connect(self.check_rarely_used_payment_method)
        left_layout.addWidget(btn_nadir_odeme)
        
        btn_birim_fiyat = QPushButton("📈 Birim Fiyat Artışı")
        btn_birim_fiyat.clicked.connect(self.check_unit_price_increase)
        left_layout.addWidget(btn_birim_fiyat)
        
        left_layout.addSpacing(10)
        left_layout.addWidget(QLabel("Gider ve Değer Kontrolleri:"))
        
        btn_kdv = QPushButton("💸 KDV Tutarlılık")
        btn_kdv.clicked.connect(self.check_kdv_consistency)
        left_layout.addWidget(btn_kdv)
        
        btn_kkdf = QPushButton("💰 KKDF Kontrol")
        btn_kkdf.clicked.connect(self.check_kkdf_kontrol)
        left_layout.addWidget(btn_kkdf)
        
        btn_gozetim = QPushButton("🔍 Gözetim Kontrol")
        btn_gozetim.clicked.connect(self.check_gozetim_kontrol)
        left_layout.addWidget(btn_gozetim)
        
        btn_yurt_ici = QPushButton("🏠 Yurt İçi Gider Kontrol")
        btn_yurt_ici.clicked.connect(self.check_domestic_expense_variation)
        left_layout.addWidget(btn_yurt_ici)
        
        btn_yurt_disi = QPushButton("🌏 Yurt Dışı Gider Kontrol")
        btn_yurt_disi.clicked.connect(self.check_foreign_expense_variation)
        left_layout.addWidget(btn_yurt_disi)
        
        btn_supalan = QPushButton("🚢 Supalan Depolama Kontrol")
        btn_supalan.clicked.connect(self.check_supalan_storage)
        left_layout.addWidget(btn_supalan)
        
        btn_igv = QPushButton("💰 IGV Kontrol")
        btn_igv.clicked.connect(self.check_igv_consistency)
        left_layout.addWidget(btn_igv)
        
        btn_tedarikci_beyan = QPushButton("📋 Tedarikçi Beyan Kontrol")
        btn_tedarikci_beyan.clicked.connect(self.check_tedarikci_beyan_kontrol)
        left_layout.addWidget(btn_tedarikci_beyan)
        
        left_layout.addStretch()
        
        self.right_panel = QFrame()
        right_layout = QVBoxLayout(self.right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(0)
        
        self.check_results_widget = CheckResultsWidget()
        right_layout.addWidget(self.check_results_widget)
        
        analysis_layout.addWidget(self.left_panel)
        analysis_layout.addWidget(self.right_panel)
        self.tabs.addTab(self.analysis_tab, "Analizler")
        
        # Örnekleme Sekmesi (Yeni)
        self.sampling_tab = QWidget()
        sampling_layout = QVBoxLayout(self.sampling_tab)
        sampling_layout.setContentsMargins(10, 10, 10, 10)
        
        # Kontrol paneli
        sampling_control = QWidget()
        sampling_control.setObjectName("samplingPanel")
        sampling_control.setStyleSheet("""
            #samplingPanel {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(255, 255, 255, 0.98), stop:1 rgba(248, 250, 251, 0.95));
                border: 1px solid rgba(226, 232, 240, 0.8);
                border-radius: 16px;
                padding: 20px;
                margin: 10px;
            }
        """)
        sampling_control_layout = QVBoxLayout(sampling_control)
        sampling_control_layout.setContentsMargins(10, 10, 10, 10)
        sampling_control_layout.setSpacing(15)
        
        # Başlık ve açıklama
        header_layout = QHBoxLayout()
        heading = QLabel("Beyanname Örnekleme")
        heading.setStyleSheet("font-size: 16px; font-weight: bold; color: #333;")
        header_layout.addWidget(heading)
        header_layout.addStretch()
        
        # Örnekleme oranı seçimi
        sample_rate_layout = QHBoxLayout()
        sample_rate_label = QLabel("Örnekleme Oranı (%):")
        sample_rate_label.setStyleSheet("font-weight: bold;")
        self.sample_rate_combo = QComboBox()
        self.sample_rate_combo.addItems(["5", "10", "15", "20", "25"])
        self.sample_rate_combo.setCurrentIndex(0)
        sample_rate_layout.addWidget(sample_rate_label)
        sample_rate_layout.addWidget(self.sample_rate_combo)
        sample_rate_layout.addStretch()
        
        # Min-max ayarları
        min_max_layout = QHBoxLayout()
        min_label = QLabel("Minimum Örnekleme Sayısı:")
        min_label.setStyleSheet("font-weight: bold;")
        self.min_sample_spin = QSpinBox()
        self.min_sample_spin.setRange(50, 500)
        self.min_sample_spin.setValue(100)
        self.min_sample_spin.setSingleStep(10)
        
        max_label = QLabel("Maksimum Örnekleme Sayısı:")
        max_label.setStyleSheet("font-weight: bold;")
        self.max_sample_spin = QSpinBox()
        self.max_sample_spin.setRange(100, 1000)
        self.max_sample_spin.setValue(150)
        self.max_sample_spin.setSingleStep(10)
        
        min_max_layout.addWidget(min_label)
        min_max_layout.addWidget(self.min_sample_spin)
        min_max_layout.addSpacing(20)
        min_max_layout.addWidget(max_label)
        min_max_layout.addWidget(self.max_sample_spin)
        min_max_layout.addStretch()
        
        # Butonlar
        buttons_layout = QHBoxLayout()
        
        # Örnekleme başlat butonu
        self.start_sampling_btn = QPushButton("🚀 Örnekleme Başlat")
        self.start_sampling_btn.setProperty("class", "success")
        self.start_sampling_btn.clicked.connect(self.start_sampling)
        
        # Excel'e aktar butonu
        self.export_excel_btn = QPushButton("📊 Excel'e Aktar")
        self.export_excel_btn.setProperty("class", "primary")
        self.export_excel_btn.clicked.connect(self.export_sampling_to_excel)
        self.export_excel_btn.setEnabled(False)  # Başlangıçta devre dışı
        
        # Temizle butonu
        self.clear_sampling_btn = QPushButton("🧹 Temizle")
        self.clear_sampling_btn.setProperty("class", "danger")
        self.clear_sampling_btn.clicked.connect(self.clear_sampling)
        self.clear_sampling_btn.setEnabled(False)  # Başlangıçta devre dışı
        
        buttons_layout.addWidget(self.start_sampling_btn)
        buttons_layout.addWidget(self.export_excel_btn)
        buttons_layout.addWidget(self.clear_sampling_btn)
        buttons_layout.addStretch()
        
        # Kontrol paneline tüm layout'ları ekle
        sampling_control_layout.addLayout(header_layout)
        sampling_control_layout.addWidget(QLabel("Bu modül, ithalat beyannamelerinden belirli kriterlere göre örnekleme yaparak faaliyet raporuna dahil edilecek beyannameleri seçer."))
        sampling_control_layout.addLayout(sample_rate_layout)
        sampling_control_layout.addLayout(min_max_layout)
        sampling_control_layout.addLayout(buttons_layout)
        
        # Kontrol panelini ana layout'a ekle
        sampling_layout.addWidget(sampling_control)
        
        # Sonuç gösterimi için tablo
        self.sampling_results_label = QLabel("Örnekleme Sonuçları")
        self.sampling_results_label.setStyleSheet("font-size: 14px; font-weight: bold; color: #333;")
        sampling_layout.addWidget(self.sampling_results_label)
        
        # Örnekleme sonuçları viewer'ı
        self.sampling_viewer = DataFrameViewer()
        sampling_layout.addWidget(self.sampling_viewer)
        
        # İstatistik bilgileri
        self.sampling_stats_label = QLabel("Örnekleme henüz yapılmadı")
        self.sampling_stats_label.setStyleSheet("color: #757575; font-style: italic;")
        sampling_layout.addWidget(self.sampling_stats_label)
        
        # Sekmeyi ana tab widget'a ekle
        self.tabs.addTab(self.sampling_tab, "Örnekleme")
        
        # Dashboard tab - Simple initialization
        self.dashboard_tab = QWidget()
        self.tabs.addTab(self.dashboard_tab, "Dashboard")
        
        # Açılışta ilk sekme "Veri Görünümü" olarak ayarla
        self.tabs.setCurrentIndex(0)
        
        # Tab değişimi için event handler ekle
        self.tabs.currentChanged.connect(self.on_tab_changed)
        
        main_layout.addWidget(self.tabs)
        
        # QtWebEngine önbellek hatalarını önle
        self.configure_qt_web_engine()
    
    def _dashboard_card(self, title, value, color="#4299e1", icon="📊"):
        """Dashboard için modern istatistik kartı oluştur"""
        w = QWidget()
        w.setStyleSheet(f"""
            QWidget {{
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(255, 255, 255, 0.95), stop:1 rgba(248, 250, 251, 0.9));
                border: 1px solid rgba(226, 232, 240, 0.8);
                border-left: 4px solid {color};
                border-radius: 12px;
                padding: 20px;
                margin: 8px;
            }}
        """)
        w.setMinimumWidth(200)
        w.setMaximumWidth(280)
        w.setMinimumHeight(120)
        w.setMaximumHeight(140)
        
        layout = QVBoxLayout(w)
        layout.setContentsMargins(15, 12, 15, 12)
        layout.setSpacing(8)
        
        # İkon ve başlık satırı
        header_layout = QHBoxLayout()
        header_layout.setSpacing(8)
        
        # İkon
        icon_label = QLabel(icon)
        icon_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                background: rgba(66, 153, 225, 0.1);
                border-radius: 8px;
                padding: 8px;
                min-width: 20px;
                max-width: 40px;
                min-height: 20px;
                max-height: 40px;
            }
        """)
        icon_label.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(icon_label)
        
        # Başlık
        title_label = QLabel(title)
        title_label.setStyleSheet("""
            QLabel {
                font-size: 12px;
                font-weight: 600;
                color: #4a5568;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                background: transparent;
            }
        """)
        title_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        title_label.setWordWrap(True)
        header_layout.addWidget(title_label, 1)
        
        layout.addLayout(header_layout)
        
        # Değer
        value_label = QLabel(str(value))
        value_label.setStyleSheet(f"""
            QLabel {{
                font-size: 28px;
                font-weight: 700;
                color: #1a202c;
                background: transparent;
                margin: 8px 0px;
            }}
        """)
        value_label.setAlignment(Qt.AlignLeft | Qt.AlignBottom)
        layout.addWidget(value_label)
        
        # Alt boşluk
        layout.addStretch()
        
        return w
    
    def _dashboard_summary_table(self, df):
        """Dashboard için özet tablo oluştur"""
        try:
            # En çok kullanılan 5 GTIP, ülke, rejim ve toplam tutar
            gtip_col = next((c for c in df.columns if c.lower().startswith("gtip")), None)
            ulke_col = next((c for c in df.columns if "ulke" in c.lower()), None)
            rejim_col = next((c for c in df.columns if "rejim" in c.lower()), None)
            tutar_col = next((c for c in df.columns if "tutar" in c.lower() or "fatura_miktari" in c.lower()), None)
            
            # Tablo oluştur
            table = QTableWidget(5, 4)
            table.setHorizontalHeaderLabels(["GTIP", "Ülke", "Rejim", "Toplam Tutar"])
            
            # GTIP sütunu
            if gtip_col and df[gtip_col].nunique() > 0:
                gtip_counts = df[gtip_col].value_counts()
                for i in range(min(5, len(gtip_counts))):
                    gtip_val = str(gtip_counts.index[i])[:15] + "..." if len(str(gtip_counts.index[i])) > 15 else str(gtip_counts.index[i])
                    table.setItem(i, 0, QTableWidgetItem(gtip_val))
            
            # Ülke sütunu
            if ulke_col and df[ulke_col].nunique() > 0:
                ulke_counts = df[ulke_col].value_counts()
                for i in range(min(5, len(ulke_counts))):
                    ulke_val = str(ulke_counts.index[i])
                    table.setItem(i, 1, QTableWidgetItem(ulke_val))
            
            # Rejim sütunu
            if rejim_col and df[rejim_col].nunique() > 0:
                rejim_counts = df[rejim_col].value_counts()
                for i in range(min(5, len(rejim_counts))):
                    rejim_val = str(rejim_counts.index[i])
                    table.setItem(i, 2, QTableWidgetItem(rejim_val))
            
            # Tutar sütunu
            if tutar_col and gtip_col and df[tutar_col].notna().any():
                try:
                    # Numerik tutarları dönüştür
                    df_temp = df.copy()
                    df_temp[tutar_col] = pd.to_numeric(df_temp[tutar_col], errors='coerce')
                    tutar_by_gtip = df_temp.groupby(gtip_col)[tutar_col].sum().sort_values(ascending=False)
                    
                    for i in range(min(5, len(tutar_by_gtip))):
                        tutar = f"{tutar_by_gtip.values[i]:,.0f}" if tutar_by_gtip.values[i] == tutar_by_gtip.values[i] else "0"
                        table.setItem(i, 3, QTableWidgetItem(tutar))
                except Exception:
                    # Tutar hesaplanamıyorsa boş bırak
                    pass
            
            # Boş hücreleri doldur
            for i in range(5):
                for j in range(4):
                    if table.item(i, j) is None:
                        table.setItem(i, j, QTableWidgetItem("-"))
            
            # Tablo ayarları
            table.setEditTriggers(QTableWidget.NoEditTriggers)
            table.setStyleSheet("""
                QTableWidget {
                    background: #fff; 
                    border-radius: 6px; 
                    font-size: 11px;
                    border: 1px solid #ddd;
                }
                QHeaderView::section {
                    background-color: #f5f5f5;
                    border: 1px solid #ddd;
                    padding: 5px;
                    font-weight: bold;
                }
            """)
            table.resizeColumnsToContents()
            table.setMaximumHeight(200)
            
            return table
            
        except Exception as e:
            print(f"Dashboard özet tablo hatası: {e}")
            # Hata durumunda basit bir mesaj döndür
            error_label = QLabel("Özet tablo oluşturulamadı")
            error_label.setStyleSheet("color: #999; font-style: italic;")
            return error_label
    
    def configure_qt_web_engine(self):
        """QtWebEngine kullanımı için yapılandırma"""
        try:
            from PyQt5.QtWebEngineWidgets import QWebEngineProfile
            
            # Ortak bir önbellek klasörü kullan
            profile = QWebEngineProfile.defaultProfile()
            app_data_dir = os.path.join(os.path.expanduser("~"), "AppData", "Local", "CustomsCheck")
            
            # Klasörü oluştur (yoksa)
            os.makedirs(app_data_dir, exist_ok=True)
            
            # Önbellek ve HTTP önbelleği için alt klasörler oluştur
            cache_dir = os.path.join(app_data_dir, "WebCache")
            http_cache_dir = os.path.join(app_data_dir, "HttpCache")
            
            try:
                os.makedirs(cache_dir, exist_ok=True)
                os.makedirs(http_cache_dir, exist_ok=True)
                
                # İzinleri kontrol et ve ayarla
                if not os.access(cache_dir, os.W_OK):
                    print(f"Uyarı: {cache_dir} klasörüne yazma izni yok")
                
                if not os.access(http_cache_dir, os.W_OK):
                    print(f"Uyarı: {http_cache_dir} klasörüne yazma izni yok")
            except Exception as e:
                print(f"Önbellek klasörü oluşturma hatası: {e}")
            
            # Önbellek ayarlarını yapılandır
            profile.setCachePath(cache_dir)
            profile.setHttpCachePath(http_cache_dir)
            profile.setPersistentCookiesPolicy(QWebEngineProfile.NoPersistentCookies)
            
            # Önbellek boyutunu sınırla
            profile.setHttpCacheType(QWebEngineProfile.DiskHttpCache)
            profile.setHttpCacheMaximumSize(5 * 1024 * 1024)  # 5 MB
            
            print("QtWebEngine önbellek yapılandırması tamamlandı")
            
        except ImportError:
            print("QtWebEngineWidgets modülü bulunamadı")
        except Exception as e:
            print(f"QtWebEngine yapılandırma hatası: {e}")
        
    def import_xml(self):
        """Import a single XML file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "XML Dosyası Seç", "", "XML Files (*.xml)"
        )
        
        if file_path:
            try:
                self.status_label.setText(f"İşleniyor: {os.path.basename(file_path)}")
                self.progress_bar.setVisible(True)
                self.progress_bar.setValue(50)
                
                # Process the XML file using existing function
                result = extract_beyanname_fixed(file_path)
                
                if 'dataframe' in result and result['dataframe'] is not None:
                    # Get the DataFrame directly from the result
                    df = result['dataframe']
                    
                    # Store the DataFrame
                    file_name = os.path.basename(file_path)
                    self.all_dataframes[file_name] = df
                    
                    # Update UI
                    self.update_file_selector()
                    self.file_selector.setCurrentText(file_name)
                    self.display_dataframe(df)
                    
                    self.status_label.setText(f"Yüklendi: {file_name} ({len(df)} satır, {len(df.columns)} sütun)")
                else:
                    QMessageBox.warning(self, "Hata", "DataFrame oluşturulamadı.")
                    self.status_label.setText("Hata: DataFrame oluşturulamadı")
                
                self.progress_bar.setVisible(False)
                
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"XML işleme hatası: {str(e)}")
                self.status_label.setText("Hata oluştu")
                self.progress_bar.setVisible(False)
    
    def import_xml_folder(self):
        """Import all XML files from a folder"""
        folder_path = QFileDialog.getExistingDirectory(
            self, "XML Klasörü Seç"
        )
        
        if folder_path:
            try:
                self.status_label.setText(f"Klasör işleniyor: {folder_path}")
                self.progress_bar.setVisible(True)
                self.progress_bar.setValue(0)
                QApplication.processEvents()
                
                # 1. Dosyalar yükleniyor
                self.status_label.setText("XML dosyaları yükleniyor...")
                QApplication.processEvents()
                
                def update_progress(progress, message):
                    # progress: 0-1 arası float, 0-70 arası ölçekle
                    bar_value = int(progress * 70)
                    self.progress_bar.setValue(bar_value)
                    self.status_label.setText(message)
                    QApplication.processEvents()
                
                dataframes, error_messages = process_multiple_xml_files(
                    folder_path, 
                    progress_callback=update_progress
                )
                
                self.progress_bar.setValue(75)
                self.status_label.setText("Veriler birleştiriliyor...")
                QApplication.processEvents()
                
                # 2. Kısa bir bekleme animasyonu (0.2 sn)
                import time
                time.sleep(0.2)
                self.progress_bar.setValue(85)
                QApplication.processEvents()
                
                # 3. DataFrame'leri sakla
                self.all_dataframes.update(dataframes)
                
                # 4. UI'ı güncelle
                self.update_file_selector()
                if self.file_selector.count() > 0:
                    self.file_selector.setCurrentIndex(0)
                
                processed_count = len(dataframes)
                error_count = len(error_messages)
                
                if error_count > 0:
                    message = f"{processed_count} XML dosyası başarıyla işlendi, {error_count} dosyada hata oluştu."
                else:
                    message = f"{processed_count} XML dosyası başarıyla işlendi."
                
                self.progress_bar.setValue(95)
                self.status_label.setText(message + " Birleştirilen veri hazırlanıyor...")
                QApplication.processEvents()
                
                # 5. Birleştir ve göster
                self.show_merged_dataframes()
                self.progress_bar.setValue(100)
                self.status_label.setText("Tüm dosyalar başarıyla birleştirildi ve gösteriliyor.")
                QApplication.processEvents()
                
                # 6. Progress bar'ı kısa süre sonra gizle
                from PyQt5.QtCore import QTimer
                QTimer.singleShot(800, lambda: self.progress_bar.setVisible(False))
                
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"Klasör işleme hatası: {str(e)}")
                self.status_label.setText("Hata oluştu")
                self.progress_bar.setVisible(False)
    
    def update_file_selector(self):
        """Update the file selector dropdown with loaded files"""
        self.file_selector.clear()
        self.file_selector.addItems(self.all_dataframes.keys())
    
    def change_active_file(self, index):
        """Change the active file when selected from dropdown"""
        # Dosya seçiciyi yeniden aktifleştir (birleştirme modundan çıkış)
        self.file_selector.setEnabled(True)
        
        # Pencere başlığını sıfırla
        self.setWindowTitle("Beyanname Kontrol Uygulaması")
        
        if index >= 0:
            file_name = self.file_selector.currentText()
            if file_name in self.all_dataframes:
                df = self.all_dataframes[file_name]
                self.display_dataframe(df)
    
    def display_dataframe(self, df):
        """Display a DataFrame in the data viewer and update other components"""
        self.current_df = df
        
        # Update data viewer
        self.data_viewer.set_dataframe(df)
        
        # Örnekleme aracını güncelle
        if hasattr(self, 'sampling_tool'):
            self.sampling_tool.set_dataframe(df)
            # Örnekleme daha önce yapıldıysa temizle
            if hasattr(self, 'export_excel_btn') and self.export_excel_btn.isEnabled():
                self.clear_sampling()
        
        # Update dashboard
        self.update_dashboard()
    
    def update_dashboard(self):
        """Update dashboard with current data"""
        # Dashboard tab'ının layout'unu kontrol et ve oluştur
        if self.dashboard_tab.layout() is None:
            layout = QVBoxLayout(self.dashboard_tab)
            layout.setContentsMargins(10, 10, 10, 10)
        else:
            layout = self.dashboard_tab.layout()
            # Eski widget'ları temizle
            while layout.count():
                item = layout.takeAt(0)
                widget = item.widget()
                if widget is not None:
                    widget.deleteLater()
        
        # İçeriği oluştur
        if self.current_df is None or (hasattr(self.current_df, 'empty') and self.current_df.empty):
            msg = QLabel("Dashboard için veri yüklenmedi. Lütfen bir XML dosyası yükleyin.")
            msg.setAlignment(Qt.AlignCenter)
            msg.setStyleSheet("font-size: 15px; color: #888; margin-top: 60px;")
            layout.addStretch()
            layout.addWidget(msg)
            layout.addStretch()
        else:
            df = self.current_df
            
            # Scroll area oluştur (çok fazla içerik olacak)
            scroll_area = QScrollArea()
            scroll_area.setWidgetResizable(True)
            scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            
            # Scroll içeriği
            scroll_content = QWidget()
            scroll_layout = QVBoxLayout(scroll_content)
            scroll_layout.setSpacing(15)
            scroll_layout.setContentsMargins(10, 10, 10, 10)
            
            # Başlık
            title_label = QLabel("🏛️ GÜMRİK YÖNETİM PANELİ")
            title_label.setStyleSheet("""
                font-size: 24px; 
                font-weight: bold; 
                color: #2c3e50; 
                margin-bottom: 20px; 
                text-align: center;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                    stop:0 #74b9ff, stop:1 #0984e3);
                color: white;
                padding: 15px;
                border-radius: 10px;
            """)
            title_label.setAlignment(Qt.AlignCenter)
            scroll_layout.addWidget(title_label)
            
            # İlk satır - Ana göstergeler
            row1 = QWidget()
            row1_layout = QGridLayout(row1)
            row1_layout.setSpacing(15)
            
            # Temel istatistikler hesapla
            total_rows = len(df)
            total_cols = len(df.columns)
            missing_values = df.isnull().sum().sum()
            
            # Beyanname sayısı
            beyanname_col = next((c for c in df.columns if "beyanname" in c.lower() and "no" in c.lower()), None)
            unique_beyannames = df[beyanname_col].nunique() if beyanname_col else total_rows
            
            # İstatistiki kıymet hesapla
            kiymet_cols = [c for c in df.columns if any(keyword in c.lower() for keyword in ["kiymet", "value", "fatura_miktari", "tutar"])]
            total_value = 0
            avg_value = 0
            max_value = 0
            if kiymet_cols:
                try:
                    for col in kiymet_cols:
                        numeric_vals = pd.to_numeric(df[col], errors='coerce').fillna(0)
                        total_value += numeric_vals.sum()
                    avg_value = total_value / unique_beyannames if unique_beyannames > 0 else 0
                    max_value = max([pd.to_numeric(df[col], errors='coerce').max() for col in kiymet_cols if not pd.to_numeric(df[col], errors='coerce').empty])
                except:
                    pass
            
            # Ağırlık hesapla
            agirlik_cols = [c for c in df.columns if "agirlik" in c.lower()]
            total_weight = 0
            avg_weight = 0
            if agirlik_cols:
                try:
                    for col in agirlik_cols:
                        numeric_vals = pd.to_numeric(df[col], errors='coerce').fillna(0)
                        total_weight += numeric_vals.sum()
                    avg_weight = total_weight / unique_beyannames if unique_beyannames > 0 else 0
                except:
                    pass
            
            # Firma sayısı
            firma_cols = [c for c in df.columns if any(keyword in c.lower() for keyword in ["firma", "ithalatci", "gonderen", "company"])]
            unique_firms = 0
            if firma_cols:
                unique_firms = df[firma_cols[0]].nunique()
            
            # GTIP sayısı
            gtip_col = next((c for c in df.columns if c.lower().startswith("gtip")), None)
            unique_gtips = df[gtip_col].nunique() if gtip_col else 0
            
            # Ülke sayısı
            ulke_col = next((c for c in df.columns if "ulke" in c.lower()), None)
            unique_countries = df[ulke_col].nunique() if ulke_col else 0
            
            # Döviz türü sayısı
            doviz_col = next((c for c in df.columns if "doviz" in c.lower()), None)
            unique_currencies = df[doviz_col].nunique() if doviz_col else 0
            
            # Kartları ekle - İlk satır (Ana göstergeler)
            row1_layout.addWidget(self._dashboard_card("TOPLAM SATIR", f"{total_rows:,}", "#4299e1", "📊"), 0, 0)
            row1_layout.addWidget(self._dashboard_card("BEYANNAME SAYISI", f"{unique_beyannames:,}", "#48bb78", "📜"), 0, 1)
            row1_layout.addWidget(self._dashboard_card("İSTATİSTİKİ KIYMET", f"₺{total_value:,.0f}", "#9f7aea", "💰"), 0, 2)
            row1_layout.addWidget(self._dashboard_card("TOPLAM AĞIRLIK", f"{total_weight:,.0f} KG", "#ed8936", "⚖️"), 0, 3)
            row1_layout.addWidget(self._dashboard_card("FIRMA SAYISI", f"{unique_firms:,}", "#38b2ac", "🏢"), 0, 4)
            
            scroll_layout.addWidget(row1)
            
            # İkinci satır - Ortalama ve detay göstergeler
            row2 = QWidget()
            row2_layout = QGridLayout(row2)
            row2_layout.setSpacing(15)
            
            row2_layout.addWidget(self._dashboard_card("ORTALAMA KIYMET", f"₺{avg_value:,.0f}", "#38b2ac", "📊"), 0, 0)
            row2_layout.addWidget(self._dashboard_card("MAX KIYMET", f"₺{max_value:,.0f}", "#f56565", "📈"), 0, 1)
            row2_layout.addWidget(self._dashboard_card("ORT. AĞIRLIK", f"{avg_weight:,.1f} KG", "#667eea", "⚖️"), 0, 2)
            row2_layout.addWidget(self._dashboard_card("GTİP TÜRÜ", f"{unique_gtips:,}", "#ed8936", "🏷️"), 0, 3)
            row2_layout.addWidget(self._dashboard_card("ÜLKE SAYISI", f"{unique_countries:,}", "#4299e1", "🌍"), 0, 4)
            
            scroll_layout.addWidget(row2)
            
            # Üçüncü satır - Kalite ve sistem göstergeleri
            row3 = QWidget()
            row3_layout = QGridLayout(row3)
            row3_layout.setSpacing(15)
            
            # Veri kalitesi hesapla
            completeness = ((len(df) * len(df.columns) - missing_values) / (len(df) * len(df.columns)) * 100) if len(df) > 0 else 0
            
            # KDV hesapla (varsa)
            kdv_cols = [c for c in df.columns if "kdv" in c.lower()]
            total_kdv = 0
            if kdv_cols:
                try:
                    for col in kdv_cols:
                        numeric_vals = pd.to_numeric(df[col], errors='coerce').fillna(0)
                        total_kdv += numeric_vals.sum()
                except:
                    pass
            
            # Gümrük vergisi hesapla (varsa)
            vergi_cols = [c for c in df.columns if "vergi" in c.lower() and "kdv" not in c.lower()]
            total_customs_duty = 0
            if vergi_cols:
                try:
                    for col in vergi_cols:
                        numeric_vals = pd.to_numeric(df[col], errors='coerce').fillna(0)
                        total_customs_duty += numeric_vals.sum()
                except:
                    pass
            
            row3_layout.addWidget(self._dashboard_card("VERİ TAMLIĞI", f"%{completeness:.1f}", "#48bb78", "✅"), 0, 0)
            row3_layout.addWidget(self._dashboard_card("TOPLAM KDV", f"₺{total_kdv:,.0f}", "#f56565", "💸"), 0, 1)
            row3_layout.addWidget(self._dashboard_card("GÜMRÜK VERGİSİ", f"₺{total_customs_duty:,.0f}", "#9f7aea", "🏛️"), 0, 2)
            row3_layout.addWidget(self._dashboard_card("DÖVİZ TÜRÜ", f"{unique_currencies:,}", "#ed8936", "💱"), 0, 3)
            row3_layout.addWidget(self._dashboard_card("EKSİK VERİ", f"{missing_values:,}", "#f56565", "❌"), 0, 4)
            
            scroll_layout.addWidget(row3)
            
            # Özet tablo
            try:
                summary_table = self._dashboard_summary_table(df)
                if summary_table:
                    summary_title = QLabel("📈 EN ÇOK KULLANILAN DEĞERLER")
                    summary_title.setStyleSheet("""
                        font-size: 16px; 
                        font-weight: bold; 
                        color: #2c3e50; 
                        margin: 20px 0 10px 0;
                        background: #ecf0f1;
                        padding: 10px;
                        border-radius: 5px;
                    """)
                    scroll_layout.addWidget(summary_title)
                    scroll_layout.addWidget(summary_table)
            except Exception as e:
                print(f"Özet tablo oluşturma hatası: {e}")
            
            # Grafikler bölümü
            try:
                from analysis_modules import create_bar_chart, create_pie_chart
                
                # Grafik başlığı
                chart_title = QLabel("📊 VERİ DAĞILIM GRAFİKLERİ VE ANALİZLER")
                chart_title.setStyleSheet("""
                    font-size: 16px; 
                    font-weight: bold; 
                    color: #2c3e50; 
                    margin: 20px 0 10px 0;
                    background: #3498db;
                    color: white;
                    padding: 10px;
                    border-radius: 5px;
                """)
                scroll_layout.addWidget(chart_title)
                
                # Grafik container'ı - 2 satır halinde
                charts_row1 = QWidget()
                charts_row1_layout = QHBoxLayout(charts_row1)
                charts_row1_layout.setSpacing(20)
                
                # GTIP dağılım grafiği
                if gtip_col and df[gtip_col].nunique() > 1:
                    try:
                        fig1 = create_bar_chart(df, gtip_col, title="GTIP Dağılımı (İlk 10)", limit=10)
                        if fig1:
                            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
                            canvas1 = FigureCanvas(fig1)
                            canvas1.setMinimumHeight(300)
                            canvas1.setMaximumHeight(400)
                            charts_row1_layout.addWidget(canvas1)
                    except Exception as e:
                        print(f"GTIP grafiği oluşturma hatası: {e}")
                
                # Ülke dağılım grafiği
                if ulke_col and df[ulke_col].nunique() > 1:
                    try:
                        fig2 = create_pie_chart(df, ulke_col, title="Ülke Dağılımı (İlk 5)", limit=5)
                        if fig2:
                            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
                            canvas2 = FigureCanvas(fig2)
                            canvas2.setMinimumHeight(300)
                            canvas2.setMaximumHeight(400)
                            charts_row1_layout.addWidget(canvas2)
                    except Exception as e:
                        print(f"Ülke grafiği oluşturma hatası: {e}")
                
                if charts_row1_layout.count() > 0:
                    scroll_layout.addWidget(charts_row1)
                
                # İkinci grafik satırı
                charts_row2 = QWidget()
                charts_row2_layout = QHBoxLayout(charts_row2)
                charts_row2_layout.setSpacing(20)
                
                # Rejim dağılım grafiği
                rejim_col = next((c for c in df.columns if "rejim" in c.lower()), None)
                if rejim_col and df[rejim_col].nunique() > 1:
                    try:
                        fig3 = create_pie_chart(df, rejim_col, title="Rejim Dağılımı", limit=8)
                        if fig3:
                            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
                            canvas3 = FigureCanvas(fig3)
                            canvas3.setMinimumHeight(300)
                            canvas3.setMaximumHeight(400)
                            charts_row2_layout.addWidget(canvas3)
                    except Exception as e:
                        print(f"Rejim grafiği oluşturma hatası: {e}")
                
                # Döviz dağılım grafiği
                if doviz_col and df[doviz_col].nunique() > 1:
                    try:
                        fig4 = create_pie_chart(df, doviz_col, title="Döviz Türü Dağılımı", limit=6)
                        if fig4:
                            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
                            canvas4 = FigureCanvas(fig4)
                            canvas4.setMinimumHeight(300)
                            canvas4.setMaximumHeight(400)
                            charts_row2_layout.addWidget(canvas4)
                    except Exception as e:
                        print(f"Döviz grafiği oluşturma hatası: {e}")
                
                if charts_row2_layout.count() > 0:
                    scroll_layout.addWidget(charts_row2)
                    
            except Exception as e:
                print(f"Grafik oluşturma genel hatası: {e}")
            
            # Scroll area'yı ayarla
            scroll_area.setWidget(scroll_content)
            layout.addWidget(scroll_area)
    
    # Check functions
    def run_all_checks(self):
        """Run all data checks"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        # İlerleme çubuğunu göster
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(10)
        self.status_label.setText("Tüm kontroller çalıştırılıyor...")
        QApplication.processEvents()
        
        results = {}
        
        # İlerleme bildirimi için fonksiyon
        def update_progress(value, message):
            # Değeri ölçeklendir (0-100 aralığına)
            scaled_value = int(10 + (value * 0.8))  # 10-90 arası
            self.progress_bar.setValue(scaled_value)
            self.status_label.setText(message)
            QApplication.processEvents()
        
        # GTIP-Ticari Tanım tutarlılık kontrolü
        self.status_label.setText("GTIP-Ticari Tanım kontrolü yapılıyor...")
        self.progress_bar.setValue(30)
        QApplication.processEvents()
        
        gtip_ticari_tanim_check = check_gtip_ticari_tanim_consistency(self.current_df)
        if gtip_ticari_tanim_check is not None:
            gtip_ticari_tanim_check["type"] = "gtip_ticari_tanim"  # Kontrol türünü belirt
            results["GTIP-Ticari Tanım Kontrolü"] = gtip_ticari_tanim_check
        
        # GTİP-Tanım Detay Analizi
        self.status_label.setText("GTİP-Tanım Detay Analizi yapılıyor...")
        self.progress_bar.setValue(30)
        QApplication.processEvents()
        
        # Gerekli sütunları kontrol et
        required_columns_detail = ['Gtip', 'Ticari_tanimi']
        missing_columns_detail = [col for col in required_columns_detail if col not in self.current_df.columns]
        
        if not missing_columns_detail:
            try:
                # Detay analizi çalıştır
                filtered_df = self.current_df[
                    self.current_df['Ticari_tanimi'].notna() & 
                    (self.current_df['Ticari_tanimi'] != '') &
                    self.current_df['Gtip'].notna() & 
                    (self.current_df['Gtip'] != '')
                ].copy()
                
                if len(filtered_df) > 0:
                    grouped = filtered_df.groupby('Ticari_tanimi')['Gtip'].unique().reset_index()
                    grouped['GTİP_Sayısı'] = grouped['Gtip'].apply(len)
                    multiple_gtips = grouped[grouped['GTİP_Sayısı'] > 1]
                    
                    if not multiple_gtips.empty:
                        # Detaylı sonuç oluştur (kısaltılmış versiyon)
                        result_rows = []
                        for _, row in multiple_gtips.head(10).iterrows():  # İlk 10 ticari tanım için
                            ticari_tanim = row['Ticari_tanimi']
                            ticari_rows = filtered_df[filtered_df['Ticari_tanimi'] == ticari_tanim].head(5)  # Her biri için 5 kayıt
                            
                            for _, data_row in ticari_rows.iterrows():
                                result_row = {
                                    'Ticari_tanimi': ticari_tanim,
                                    'Gtip': data_row.get('Gtip', ''),
                                    'Adi_unvani': data_row.get('Adi_unvani', ''),
                                    'Beyanname_no': data_row.get('Beyanname_no', ''),
                                }
                                result_rows.append(result_row)
                        
                        result_df = pd.DataFrame(result_rows)
                        
                        results["GTİP-Tanım Detay Analizi"] = {
                            "status": "warning",
                            "message": f"{len(multiple_gtips)} ticari tanımda farklı GTİP kodları tespit edildi. (Özet: ilk 10 tanım)",
                            "data": result_df,
                            "type": "gtip_tanim_detail"
                        }
                    else:
                        results["GTİP-Tanım Detay Analizi"] = {
                            "status": "ok",
                            "message": "Aynı ticari tanımda farklı GTİP kodu kullanımı tespit edilmedi.",
                            "type": "gtip_tanim_detail"
                        }
                else:
                    results["GTİP-Tanım Detay Analizi"] = {
                        "status": "warning",
                        "message": "Analiz için uygun veri bulunamadı.",
                        "type": "gtip_tanim_detail"
                    }
            except Exception as e:
                results["GTİP-Tanım Detay Analizi"] = {
                    "status": "error",
                    "message": f"Analiz sırasında hata: {str(e)}",
                    "type": "gtip_tanim_detail"
                }
        else:
            missing_cols_str = ", ".join(missing_columns_detail)
            results["GTİP-Tanım Detay Analizi"] = {
                "status": "error",
                "message": f"Kontrol için gerekli sütunlar eksik: {missing_cols_str}",
                "type": "gtip_tanim_detail"
            }
        
        # Alıcı-Satıcı ilişki kontrolü (firma seçimi olmadan çalıştır)
        self.status_label.setText("Alıcı-Satıcı ilişki kontrolü yapılıyor...")
        self.progress_bar.setValue(30)
        QApplication.processEvents()
        
        if "Alici_satici_iliskisi" in self.current_df.columns:
            alici_satici_check = check_alici_satici_relationship(self.current_df, 
                                                              progress_callback=update_progress)
            if alici_satici_check is not None:
                alici_satici_check["type"] = "alici_satici_relationship"  # Kontrol türünü belirt
                results["Alıcı-Satıcı İlişki Kontrolü"] = alici_satici_check
        
        # İşlem Niteliği kontrolü
        self.status_label.setText("İşlem Niteliği kontrolü yapılıyor...")
        self.progress_bar.setValue(60)
        QApplication.processEvents()
        
        # İşlem Niteliği kontrolü için gerekli sütunların varlığını kontrol et
        required_columns = ['Kalem_Islem_Niteligi', 'Odeme_sekli', 'Rejim']
        missing_columns = [col for col in required_columns if col not in self.current_df.columns]
        
        if not missing_columns:
            # 1. Kontrol: Ödeme şekli "bedelsiz" ise işlem niteliği kodu "99" olmalı
            bedelsiz_payment_filter = self.current_df['Odeme_sekli'].str.lower().str.contains('bedelsiz', na=False)
            incorrect_payment_code = self.current_df[bedelsiz_payment_filter & (self.current_df['Kalem_Islem_Niteligi'] != '99')]
            
            # 2. Kontrol: Rejim kodu "6123" ise işlem niteliği kodu "61" olmalı
            rejim_filter = self.current_df['Rejim'] == '6123'
            incorrect_rejim_code = self.current_df[rejim_filter & (self.current_df['Kalem_Islem_Niteligi'] != '61')]
            
            # Tüm tutarsızlıkları birleştir
            all_inconsistencies = pd.concat([incorrect_payment_code, incorrect_rejim_code]).drop_duplicates()
            
            # Sonuçları hazırla
            if len(all_inconsistencies) > 0:
                results["İşlem Niteliği Kontrolü"] = {
                    "status": "warning",
                    "message": f"{len(all_inconsistencies)} adet tutarsız işlem niteliği kodu bulundu.",
                    "data": all_inconsistencies,
                    "type": "islem_niteligi_consistency"
                }
            else:
                results["İşlem Niteliği Kontrolü"] = {
                    "status": "ok",
                    "message": "Tüm işlem niteliği kodları ödeme şekli ve rejim kodu ile tutarlı.",
                    "type": "islem_niteligi_consistency"
                }
        else:
            missing_cols_str = ", ".join(missing_columns)
            results["İşlem Niteliği Kontrolü"] = {
                "status": "error",
                "message": f"Kontrol için gerekli sütunlar eksik: {missing_cols_str}",
                "type": "islem_niteligi_consistency"
            }
        
        # Update the check results widget
        self.progress_bar.setValue(80)
        self.status_label.setText("Sonuçlar gösteriliyor...")
        QApplication.processEvents()
        
        # Sonuçları CheckResultsWidget'a aktar
        self.check_results_widget.set_check_results(results, self.current_df)
        
        # Otomatik olarak ilk sonucu seç
        if self.check_results_widget.results_list.count() > 0:
            self.check_results_widget.results_list.setCurrentRow(0)
            self.check_results_widget.on_result_item_clicked(self.check_results_widget.results_list.currentItem())
        
        # İlerleme çubuğunu gizle
        self.progress_bar.setValue(100)
        self.status_label.setText("Tüm kontroller tamamlandı")
        QApplication.processEvents()
        self.progress_bar.setVisible(False)
    
    def check_missing_values(self):
        """Check for missing values"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        missing = check_missing_values(self.current_df)
        if missing is not None:
            result = {
                "Eksik Değer Kontrolü": {
                    "status": "warning" if not missing.empty else "ok",
                    "message": f"{missing.shape[0]} sütunda eksik değer bulundu." if not missing.empty else "Eksik değer bulunamadı.",
                    "data": missing
                }
            }
            self.check_results_widget.set_check_results(result)
            if not missing.empty:
                self.check_results_widget.show_details(missing)
        else:
            result = {
                "Eksik Değer Kontrolü": {
                    "status": "ok",
                    "message": "Eksik değer bulunamadı."
                }
            }
            self.check_results_widget.set_check_results(result)
    
    def check_duplicate_rows(self):
        """Check for duplicate rows"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        duplicates = check_duplicate_rows(self.current_df)
        result = {
            "Tekrarlanan Veri Kontrolü": {
                "status": "warning" if duplicates["duplicate_rows_all"] > 0 else "ok",
                "message": f"{duplicates['duplicate_rows_all']} tekrarlanan satır bulundu." if duplicates["duplicate_rows_all"] > 0 else "Tekrarlanan satır bulunamadı."
            }
        }
        self.check_results_widget.set_check_results(result)
        
        if duplicates["duplicate_rows_all"] > 0:
            # Show the duplicate rows
            duplicate_data = self.current_df[self.current_df.duplicated(keep='first')]
            self.check_results_widget.show_details(duplicate_data)
    
    def check_weight_consistency(self):
        """Check for weight consistency (Brut_agirlik >= Net_agirlik)"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        weight_check = check_weight_consistency(self.current_df)
        if weight_check is not None:
            result = {"Ağırlık Kontrolü": weight_check}
            self.check_results_widget.set_check_results(result)
            
            if weight_check["status"] == "warning" and "inconsistent_rows" in weight_check:
                self.check_results_widget.show_details(weight_check["inconsistent_rows"])
        else:
            result = {
                "Ağırlık Kontrolü": {
                    "status": "ok",
                    "message": "Ağırlık verisi kontrol edilemedi."
                }
            }
            self.check_results_widget.set_check_results(result)
    
    def check_islem_niteligi_consistency(self):
        """Check for Kalem_Islem_Niteligi consistency with payment method and regime code"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("İşlem Niteliği tutarlılık kontrolü çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()  # UI güncellemesi için
            
            # Modülü kullanarak kontrol yap
            self.progress_bar.setValue(50)
            QApplication.processEvents()
            
            result = kontrol_islem_niteligi_tutarlilik(self.current_df)
            
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # Sonuçları işle
            if result["status"] == "warning":
                # Hata bulundu
                check_result = {
                    "İşlem Niteliği Kontrolü": {
                        "status": "warning",
                        "message": result["message"],
                        "data": result["data"],
                        "type": "islem_niteligi_consistency",
                        "html_report": result["html_report"]
                    }
                }
                
                if "summary" in result:
                    check_result["İşlem Niteliği Kontrolü"]["summary"] = result["summary"]
                
                # Widget'a sonuçları gönder
                self.check_results_widget.set_check_results(check_result, self.current_df)
                
                # Tabloyu göster
                try:
                    self.check_results_widget.show_details(result["data"])
                    if "summary" in result:
                        self.check_results_widget.show_summary(result["summary"])
                except Exception as e:
                    print(f"Tablo gösterilirken hata: {str(e)}")
                
                self.status_label.setText(f"İşlem Niteliği Kontrolü: {result['message']}")
                
            elif result["status"] == "ok":
                # Kontrol başarılı
                check_result = {
                    "İşlem Niteliği Kontrolü": {
                        "status": "ok",
                        "message": result["message"],
                        "html_report": result["html_report"]
                    }
                }
                self.check_results_widget.set_check_results(check_result)
                self.status_label.setText("İşlem Niteliği Kontrolü: Tutarlılık kontrolü başarılı. Tüm veriler uyumlu.")
                
            else:
                # Hata durumu
                QMessageBox.critical(self, "Hata", result["message"])
                self.status_label.setText("İşlem Niteliği kontrolünde hata oluştu")
                
            self.progress_bar.setVisible(False)
            
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"İşlem Niteliği kontrolü sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)
    
    def check_gtip_ticari_tanim_consistency(self):
        """Check for GTIP-Ticari Tanım consistency - REMOVED BY USER REQUEST"""
        QMessageBox.information(self, "Bilgi", "Bu özellik kaldırılmıştır.")
        return
    
    def check_alici_satici_relationship(self):
        """Alıcı-Satıcı ilişki kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # Gönderici sütununu belirle
            sender_column = None
            possible_sender_columns = ["Gonderen", "Gonderen_adi", "Gonderen_firma", "Adi_unvani", "Ihracatci"]
            
            for col in possible_sender_columns:
                if col in self.current_df.columns:
                    sender_column = col
                    break
            
            # İşleme başlamadan kullanıcıya bilgilendir
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Alıcı-Satıcı ilişki kontrolü başlatılıyor.")
            msgBox.setInformativeText("İki farklı kontrol yöntemi vardır:\n\n"
                                    "1. Belirli firmaların ilişki durumu 6 olan beyannamelerini bulma\n"
                                    "2. Aynı göndericide hem 6 hem 0 ilişki durumu olan beyannameleri bulma\n\n"
                                    "Firma seçmek ister misiniz?")
            msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
            msgBox.setDefaultButton(QMessageBox.Yes)
            response = msgBox.exec()
            
            if response == QMessageBox.Cancel:
                return
        
            # İşlem başladığını göster
            self.status_label.setText("Alıcı-Satıcı ilişki kontrolü çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            selected_companies = None
            
            # Callback fonksiyon oluştur
            def update_progress(value, message):
                self.progress_bar.setValue(value)
                self.status_label.setText(message)
                QApplication.processEvents()  # UI'nin yanıt vermesini sağla
            
            # Kullanıcı firma seçmek isterse
            if response == QMessageBox.Yes:
                if not sender_column:
                    QMessageBox.warning(self, "Uyarı", "Gönderici sütunu bulunamadı.")
                    self.progress_bar.setVisible(False)
                    return
        
                update_progress(15, "Firma listesi hazırlanıyor...")
                
                # Benzersiz firma listesini al
                unique_companies = sorted(self.current_df[sender_column].dropna().unique().tolist())
                
                if not unique_companies:
                    QMessageBox.warning(self, "Uyarı", "Gönderici firma bulunamadı.")
                    self.progress_bar.setVisible(False)
                    return
                
                # Firma seçim dialog'u oluştur
                dialog = QDialog(self)
                dialog.setWindowTitle("Firma Seçimi")
                dialog.setMinimumWidth(500)
                dialog.setMinimumHeight(400)
                
                layout = QVBoxLayout(dialog)
                
                list_widget = QListWidget()
                list_widget.setSelectionMode(QAbstractItemView.MultiSelection)
                list_widget.addItems(unique_companies)
                
                layout.addWidget(QLabel("Kontrol edilecek firmaları seçin:"))
                layout.addWidget(list_widget)
                
                button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
                button_box.accepted.connect(dialog.accept)
                button_box.rejected.connect(dialog.reject)
                layout.addWidget(button_box)
                
                if dialog.exec() == QDialog.Accepted:
                    selected_items = list_widget.selectedItems()
                    if selected_items:
                        selected_companies = [item.text() for item in selected_items]
                        print(f"[DEBUG] Seçilen firmalar: {selected_companies}")
                    else:
                        # Kullanıcı firma seçmediyse iptal et
                        self.progress_bar.setVisible(False)
                        self.status_label.setText("İşlem iptal edildi")
                        return
                else:
                    # Dialog iptal edilirse işlemi durdur
                    self.progress_bar.setVisible(False)
                    self.status_label.setText("İşlem iptal edildi")
                    return
            
            # "Hayır" durumunda kullanıcıyı bilgilendir
            elif response == QMessageBox.No:
                # Büyük veri seti uyarısı
                if len(self.current_df) > 5000:
                    msg = QMessageBox()
                    msg.setIcon(QMessageBox.Warning)
                    msg.setText(f"Büyük veri seti tespit edildi ({len(self.current_df)} satır)")
                    msg.setInformativeText("Bu işlem büyük veri setlerinde zaman alabilir ve örnekleme yapılacaktır. Devam etmek istiyor musunuz?")
                    msg.setWindowTitle("Toplu XML Yükleme")
                    msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                    
                    if msg.exec() == QMessageBox.No:
                        self.progress_bar.setVisible(False)
                        self.status_label.setText("İşlem iptal edildi")
                        return
                
                update_progress(20, "Tüm firmalar için ilişki kontrolü başlatılıyor...")
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(30)
            QApplication.processEvents()
            
            # Kontrol işlemini çağır
            result = check_alici_satici_relationship(self.current_df, selected_companies, progress_callback=update_progress)
            print(f"[DEBUG] Alıcı-satıcı ilişki kontrol sonucu: {result}")
            if result and result.get('data') is not None:
                print(f"[DEBUG] Sonuç DataFrame ilk 5 satır:\n{result['data'].head()}")
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            if result["status"] == "error":
                QMessageBox.warning(self, "Uyarı", result["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(result["message"])
                return
            
            # Özet tablosu oluştur
            summary_df = pd.DataFrame()
            if result["data"] is not None:
                # Beyanname numaralarını say - tekil olarak!
                if "Beyanname_no" in result["data"].columns:
                    unique_beyannames = result["data"]["Beyanname_no"].unique()
                    beyanname_count = len(unique_beyannames)
                else:
                    beyanname_count = 0
                    
                # Firma sayısını belirle
                if sender_column in result["data"].columns:
                    unique_companies = result["data"][sender_column].unique()
                    company_count = len(unique_companies)
                else:
                    company_count = 0
                
                # Kontrol tipine göre özet tablo oluştur
                if result["type"] == "selected_companies":
                    summary_df = pd.DataFrame({
                        'Özet Bilgi': ['Seçilen Firma Sayısı', 'İlişki Durumu 6 Olan Beyanname Sayısı', 'Etkilenen Beyanname Sayısı'],
                        'Değer': [len(selected_companies) if selected_companies else 0, beyanname_count, beyanname_count]
                    })
                elif result["type"] == "all_senders_enhanced":
                    # Geliştirilmiş kontrolde toplam hata sayısını göster
                    total_beyanname_error_count = result.get("total_beyanname_error_count", beyanname_count)
                    firm_count = result.get("firm_count", company_count)
                    
                    summary_df = pd.DataFrame({
                        'Özet Bilgi': ['Tutarsızlık Bulunan Firma Sayısı', 'Hatalı Kodlu Beyanname Sayısı', 'Etkilenen Beyanname Sayısı'],
                        'Değer': [firm_count, total_beyanname_error_count, beyanname_count]
                    })
                else:
                    summary_df = pd.DataFrame({
                        'Özet Bilgi': ['Tutarsızlık Bulunan Firma Sayısı', 'Toplam Beyanname Sayısı', 'Etkilenen Beyanname Sayısı'],
                        'Değer': [company_count, beyanname_count, beyanname_count]
                    })
            
            # Sonuç ekle
            result["summary"] = summary_df
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(90)
            QApplication.processEvents()
                
            # Sonuçları göster
            if result["data"] is not None:
                # Check results widget'a sonuçları ekle
                check_result = {
                    "Alıcı-Satıcı İlişki Kontrolü": {
                        "status": result["status"],
                        "message": result["message"],
                        "data": result["data"],
                        "type": result["type"],
                        "summary": summary_df
                    }
                }
                
                # İstatistik bilgisini ekle (geliştirilmiş kontrolde)
                if "stats" in result:
                    check_result["Alıcı-Satıcı İlişki Kontrolü"]["stats"] = result["stats"]
                
                # Diğer ek bilgileri ekle
                for key in ["total_error_count", "total_beyanname_error_count", "firm_count"]:
                    if key in result:
                        check_result["Alıcı-Satıcı İlişki Kontrolü"][key] = result[key]
                
                # HTML rapor oluştur
                try:
                    html_content = self._generate_alici_satici_relationship_html(result, summary_df, sender_column)
                    check_result["Alıcı-Satıcı İlişki Kontrolü"]["html_report"] = html_content
                except Exception as e:
                    print(f"HTML rapor oluşturma hatası: {str(e)}")
                
                self.check_results_widget.set_check_results(check_result)
                
                # Özet ve detayları göster
                try:
                    self.check_results_widget.show_details(result["data"])
                    self.check_results_widget.show_summary(summary_df)
                    
                    # Uygun mesajı göster
                    if result["type"] == "selected_companies":
                        unique_beyanname_count = len(result["data"]["Beyanname_no"].unique())
                        self.status_label.setText(f"Seçilen firmalarda {unique_beyanname_count} adet ilişki durumu 6 olan beyanname bulundu.")
                    elif result["type"] == "all_senders_enhanced":
                        total_beyanname_error_count = result.get("total_beyanname_error_count", beyanname_count)
                        firm_count = result.get("firm_count", company_count)
                        self.status_label.setText(f"{firm_count} firmada toplam {total_beyanname_error_count} adet hatalı ilişki kodlu beyanname tespit edildi.")
                    else:
                        self.status_label.setText(f"{result['message']}")
                except Exception as e:
                    print(f"Detay gösterilirken hata: {str(e)}")
            else:
                # Sonuç yoksa uygun mesajı göster
                check_result = {
                    "Alıcı-Satıcı İlişki Kontrolü": {
                        "status": "ok",
                        "message": result["message"],
                        "type": result["type"],
                        "summary": summary_df
                    }
                }
                self.check_results_widget.set_check_results(check_result)
                self.status_label.setText(result["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            self.progress_bar.setVisible(False)
            QApplication.processEvents()
            
        except Exception as e:
            error_msg = f"Alıcı-Satıcı ilişki kontrolü sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def _generate_alici_satici_relationship_html(self, result, summary_df, sender_column):
        """Alıcı-Satıcı ilişki kontrolü için yönetici düzeyinde basit HTML rapor oluştur"""
        try:
            # Güvenli değerler için fallback
            if sender_column is None:
                sender_column = "Adi_unvani"  # Varsayılan sütun
            
            # Temel istatistikleri hesapla (güvenli şekilde)
            if sender_column in self.current_df.columns:
                total_companies = len(self.current_df[sender_column].dropna().unique())
            else:
                total_companies = 0
            total_transactions = len(self.current_df)
            
            # Sonuç tipine göre risk analizi
            if result.get("type") == "selected_companies":
                # Seçili firma kontrolü
                if result.get("data") is not None and not result["data"].empty:
                    # Gönderici sütunu varsa kullan, yoksa Adi_unvani kullan
                    if sender_column in result["data"].columns:
                        selected_companies_count = result["data"][sender_column].nunique()
                    elif "Adi_unvani" in result["data"].columns:
                        selected_companies_count = result["data"]["Adi_unvani"].nunique()
                        sender_column = "Adi_unvani"  # Güncelle
                    else:
                        selected_companies_count = 0
                    
                    if "Beyanname_no" in result["data"].columns:
                        problematic_beyannames = len(result["data"].drop_duplicates(subset=['Beyanname_no']))
                    else:
                        problematic_beyannames = len(result["data"])
                    
                    # Risk seviyesi belirleme (İlişki durumu 6 olanların sayısına göre)
                    if problematic_beyannames > 50:
                        overall_risk = "YÜKSEK"
                        risk_color = "#e74c3c"
                        risk_icon = "🚨"
                    elif problematic_beyannames > 10:
                        overall_risk = "ORTA"
                        risk_color = "#f39c12"
                        risk_icon = "⚠️"
                    else:
                        overall_risk = "DÜŞÜK"
                        risk_color = "#27ae60"
                        risk_icon = "✅"
                    
                    high_risk_companies = selected_companies_count if problematic_beyannames > 50 else 0
                    medium_risk_companies = selected_companies_count if 10 < problematic_beyannames <= 50 else 0
                    low_risk_companies = selected_companies_count if problematic_beyannames <= 10 else 0
                else:
                    selected_companies_count = 0
                    problematic_beyannames = 0
                    overall_risk = "DÜŞÜK"
                    risk_color = "#27ae60"
                    risk_icon = "✅"
                    high_risk_companies = 0
                    medium_risk_companies = 0
                    low_risk_companies = 0
                    
            elif result.get("type") == "all_senders_enhanced":
                # Gelişmiş analiz
                if result.get("data") is not None and "stats" in result:
                    firm_count = result.get("firm_count", 0)
                    total_beyanname_error_count = result.get("total_beyanname_error_count", 0)
                    
                    # Risk seviyesi belirleme (Hatalı beyanname oranına göre)
                    error_rate = (total_beyanname_error_count / total_transactions * 100) if total_transactions > 0 else 0
                    
                    if error_rate > 5:
                        overall_risk = "YÜKSEK"
                        risk_color = "#e74c3c"
                        risk_icon = "🚨"
                        high_risk_companies = firm_count
                        medium_risk_companies = 0
                        low_risk_companies = 0
                    elif error_rate > 1:
                        overall_risk = "ORTA"
                        risk_color = "#f39c12"
                        risk_icon = "⚠️"
                        high_risk_companies = 0
                        medium_risk_companies = firm_count
                        low_risk_companies = 0
                    else:
                        overall_risk = "DÜŞÜK"
                        risk_color = "#27ae60"
                        risk_icon = "✅"
                        high_risk_companies = 0
                        medium_risk_companies = 0
                        low_risk_companies = firm_count
                        
                    problematic_beyannames = total_beyanname_error_count
                else:
                    firm_count = 0
                    problematic_beyannames = 0
                    overall_risk = "DÜŞÜK"
                    risk_color = "#27ae60"
                    risk_icon = "✅"
                    high_risk_companies = 0
                    medium_risk_companies = 0
                    low_risk_companies = 0
            else:
                # Diğer kontrol türleri
                problematic_beyannames = 0
                overall_risk = "DÜŞÜK"
                risk_color = "#27ae60"
                risk_icon = "✅"
                high_risk_companies = 0
                medium_risk_companies = 0
                low_risk_companies = 0
            
            # Risk dağılımı yüzdeleri
            total_risk_companies = high_risk_companies + medium_risk_companies + low_risk_companies
            if total_risk_companies > 0:
                high_risk_pct = (high_risk_companies / total_risk_companies * 100)
                medium_risk_pct = (medium_risk_companies / total_risk_companies * 100)
                low_risk_pct = (low_risk_companies / total_risk_companies * 100)
            else:
                high_risk_pct = medium_risk_pct = low_risk_pct = 0
            
            # Yönetici özeti HTML
            html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Alıcı-Satıcı Risk Analizi</title>
            <style>
                    body {{
                        font-family: 'Segoe UI', Arial, sans-serif;
                margin: 0;
                        padding: 20px;
                        background: #f8f9fa;
                    }}
                    .container {{
                        max-width: 1200px;
                        margin: 0 auto;
                        background: white;
                        border-radius: 10px;
                        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
                        overflow: hidden;
                    }}
                    .header {{
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                        padding: 30px;
                        text-align: center;
                    }}
                    .header h1 {{
                        margin: 0;
                        font-size: 24px;
                        font-weight: 300;
                    }}
                    .executive-summary {{
                        padding: 30px;
                        background: #f8f9fa;
                        border-bottom: 3px solid #dee2e6;
                    }}
                    .risk-banner {{
                        background: {risk_color};
                        color: white;
                        padding: 20px;
                        text-align: center;
                        font-size: 20px;
                        font-weight: bold;
            margin-bottom: 20px;
                        border-radius: 8px;
                    }}
                    .metrics-grid {{
                        display: grid;
                        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                        gap: 20px;
                        margin: 20px 0;
                    }}
                    .metric-card {{
                        background: white;
                        padding: 20px;
                        border-radius: 8px;
                        text-align: center;
                        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                        border-left: 4px solid #3498db;
                    }}
                    .metric-value {{
                        font-size: 32px;
                        font-weight: bold;
                        color: #2c3e50;
                        margin-bottom: 5px;
                    }}
                    .metric-label {{
                        color: #7f8c8d;
                        font-size: 14px;
                        text-transform: uppercase;
                        letter-spacing: 1px;
                    }}
                    .action-required {{
                        background: #fff3cd;
                        border: 1px solid #ffeaa7;
                        border-radius: 8px;
                        padding: 20px;
                        margin: 20px 0;
                    }}
                    .action-required h3 {{
                        color: #856404;
                        margin-top: 0;
                    }}
                    .summary-table {{
                width: 100%;
                border-collapse: collapse;
                        margin: 20px 0;
                    }}
                    .summary-table th {{
                        background: #f8f9fa;
                        padding: 12px;
                text-align: left;
                        border-bottom: 2px solid #dee2e6;
                        font-weight: 600;
                    }}
                    .summary-table td {{
                padding: 12px;
                        border-bottom: 1px solid #dee2e6;
                    }}
                    .summary-table tr:hover {{
                        background: #f8f9fa;
                    }}
                    .recommendation {{
                        background: #e3f2fd;
                        border-left: 4px solid #2196f3;
                        padding: 20px;
                        margin: 20px 0;
                        border-radius: 0 8px 8px 0;
                    }}
                    .detail-table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 20px 0;
                        font-size: 12px;
                    }}
                    .detail-table th {{
                        background: linear-gradient(135deg, #1565c0 0%, #1976d2 100%);
                        color: white;
                        padding: 10px 8px;
                        text-align: left;
                        font-weight: 600;
                    }}
                    .detail-table td {{
                        padding: 8px;
                        border-bottom: 1px solid #e0e0e0;
                        vertical-align: top;
                    }}
                    .detail-table tr:nth-child(even) {{
                        background-color: #f8f9fa;
                    }}
                    .detail-table tr:hover {{
                        background-color: #e3f2fd;
                    }}
                    .beyanname-badge {{
                        background: #17a2b8;
                        color: white;
                        padding: 3px 8px;
                        border-radius: 12px;
                        font-size: 10px;
                        font-weight: 500;
                        display: inline-block;
                    }}
                    .firma-name {{
                        font-weight: 500;
                        color: #424242;
                        max-width: 200px;
                        word-wrap: break-word;
                    }}
            </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1>🏢 Alıcı-Satıcı İlişki Kontrolü</h1>
                        <p>Yönetici Özeti</p>
                    </div>
                    
                    <div class="executive-summary">
                        <div class="risk-banner">
                            {risk_icon} RİSK SEVİYESİ: {overall_risk}
                        </div>
                        
                        <div class="metrics-grid">
                            <div class="metric-card">
                                <div class="metric-value">{total_companies:,}</div>
                                <div class="metric-label">Toplam Şirket</div>
                    </div>
                            <div class="metric-card">
                                <div class="metric-value">{total_transactions:,}</div>
                                <div class="metric-label">Toplam İşlem</div>
                </div>
                            <div class="metric-card">
                                <div class="metric-value">{problematic_beyannames:,}</div>
                                <div class="metric-label">Sorunlu Beyanname</div>
                    </div>
                            <div class="metric-card">
                                <div class="metric-value">{high_risk_companies + medium_risk_companies:,}</div>
                                <div class="metric-label">İncelenmesi Gereken Şirket</div>
                        </div>
                </div>
                """
            
            # Kontrol türüne göre ek bilgi ekle
            if result.get("type") == "selected_companies":
                html += f"""
                        <div class="action-required">
                            <h3>📋 Seçili Firma Kontrolü</h3>
                            <p>Bu analiz seçilen firmaların <strong>ilişki durumu 6</strong> olan beyannamelerini listeler.</p>
                            <ul>
                                <li><strong>İlişki Durumu 6:</strong> Alıcı ve satıcı arasında dolaylı ya da dolaysız ilişki vardır</li>
                                <li>Bu beyannameler detaylı inceleme gerektirebilir</li>
                            </ul>
                    </div>
                """
                
                # Seçili firmalar için detaylı tablo ekle
                if result.get("data") is not None and not result["data"].empty:
                    data_df = result["data"]
                    
                    # Firma bazında grupla
                    if sender_column in data_df.columns:
                        firma_summary = data_df.groupby(sender_column).agg({
                            'Beyanname_no': 'nunique',
                            'Alici_satici_iliskisi': 'first'
                        }).reset_index()
                        firma_summary.columns = ['Firma', 'Beyanname_Sayisi', 'İlişki_Durumu']
                        firma_summary = firma_summary.sort_values('Beyanname_Sayisi', ascending=False)
                        
                        html += """
                                <h3 style="color: #2c3e50; margin: 20px 0;">📊 Seçili Firmaların Detayları</h3>
                                <table class="detail-table">
                                    <thead>
                                        <tr>
                                            <th style="width: 50%;">Firma Unvanı</th>
                                            <th style="width: 25%;">İlişki Durumu 6 Beyanname Sayısı</th>
                                            <th style="width: 25%;">İlişki Durumu</th>
                                        </tr>
                                    </thead>
                                    <tbody>
                        """
                        
                        for _, row in firma_summary.iterrows():
                            firma = row['Firma']
                            beyanname_sayisi = row['Beyanname_Sayisi']
                            iliski_durumu = row['İlişki_Durumu']
                            
                            html += f"""
                                        <tr>
                                            <td><div class="firma-name">{firma}</div></td>
                                            <td><span class="beyanname-badge">{beyanname_sayisi} beyanname</span></td>
                                            <td>{iliski_durumu}</td>
                                        </tr>
                            """
                        
                        html += """
                                    </tbody>
                                </table>
                        """
                        
                        # Beyanname detayları
                        html += """
                                <h3 style="color: #2c3e50; margin: 20px 0;">📋 Tüm Beyanname Detayları</h3>
                                <table class="detail-table">
                                    <thead>
                                        <tr>
                                            <th style="width: 40%;">Firma Unvanı</th>
                                            <th style="width: 30%;">Beyanname No</th>
                                            <th style="width: 15%;">İlişki Durumu</th>
                                            <th style="width: 15%;">Durum</th>
                                        </tr>
                                    </thead>
                                    <tbody>
                        """
                        
                        # Maksimum 50 beyanname göster
                        for _, row in data_df.head(50).iterrows():
                            firma = row.get(sender_column, '')
                            beyanname = row.get('Beyanname_no', '')
                            iliski_durumu = row.get('Alici_satici_iliskisi', '')
                            
                            html += f"""
                                        <tr>
                                            <td><div class="firma-name">{firma}</div></td>
                                            <td><span class="beyanname-badge">{beyanname}</span></td>
                                            <td>{iliski_durumu}</td>
                                            <td>⚠️ İnceleme Gerekli</td>
                                        </tr>
                            """
                        
                        if len(data_df) > 50:
                            html += f"""
                                        <tr>
                                            <td colspan="4" style="text-align: center; font-style: italic; padding: 15px; background-color: #f5f5f5;">
                                                <strong>📋 Toplam {len(data_df)} beyanname bulundu, ilk 50 tanesi gösterilmektedir.</strong>
                                            </td>
                                        </tr>
                            """
                        
                        html += """
                                    </tbody>
                                </table>
                        """
                        
            elif result.get("type") == "all_senders_enhanced":
                html += f"""
                        <div class="action-required">
                            <h3>🔍 Tutarsızlık Analizi</h3>
                            <p>Bu analiz <strong>aynı firmada hem ilişki durumu 0 hem de 6 olan beyannameleri</strong> tespit eder.</p>
                            <ul>
                                <li><strong>İlişki Durumu 0:</strong> Alıcı ve satıcı arasında ilişki yoktur</li>
                                <li><strong>İlişki Durumu 6:</strong> İlişki vardır ama fiyatı etkilememiştir</li>
                                <li>Aynı firmada her iki kodun kullanılması tutarsızlık gösterir</li>
                            </ul>
                </div>
                """
                
                # Tutarsız firmalar için detaylı tablo ekle
                if result.get("data") is not None and not result["data"].empty and "stats" in result:
                    stats_df = result["stats"]
                    data_df = result["data"]
                    
                    html += """
                            <h3 style="color: #2c3e50; margin: 20px 0;">📊 Tutarsız Firmaların Detayları</h3>
                            <div style="background: #e3f2fd; border: 1px solid #2196f3; border-radius: 6px; padding: 15px; margin: 15px 0;">
                                <strong>ℹ️ Bilgi:</strong> Sayılar <strong>benzersiz beyanname numarası</strong> bazındadır. 
                                Aynı beyanname numarasının farklı satırlarda görünmesi normaldir çünkü her satır ayrı kalem bilgisini temsil eder.
                            </div>
                            <table class="detail-table">
                                <thead>
                                    <tr>
                                        <th style="width: 35%;">Gönderici Firma</th>
                                        <th style="width: 15%;">İlişki Kodu 0<br><small>(Benzersiz Beyanname)</small></th>
                                        <th style="width: 15%;">İlişki Kodu 6<br><small>(Benzersiz Beyanname)</small></th>
                                        <th style="width: 15%;">Hatalı Kod</th>
                                        <th style="width: 20%;">Hatalı Beyanname Sayısı<br><small>(Benzersiz)</small></th>
                                    </tr>
                                </thead>
                                <tbody>
                    """
                    
                    for _, row in stats_df.iterrows():
                        firma = row['Firma']
                        kod_0_sayisi = row['Kod_0_Sayısı']
                        kod_6_sayisi = row['Kod_6_Sayısı']
                        hatali_kod = row['Hatalı_Kod']
                        hatali_beyanname_sayisi = row['Hatalı_Beyanname_Sayısı']
                        
                        # Hata oranına göre renk belirle
                        toplam = kod_0_sayisi + kod_6_sayisi
                        hata_orani = (hatali_beyanname_sayisi / toplam * 100) if toplam > 0 else 0
                        
                        if hata_orani > 50:
                            row_color = "#ffebee"
                            border_color = "#e74c3c"
                        elif hata_orani > 30:
                            row_color = "#fff3e0"
                            border_color = "#f39c12"
                        else:
                            row_color = "#f3e5f5"
                            border_color = "#9c27b0"
                        
                        html += f"""
                                    <tr style="background-color: {row_color}; border-left: 3px solid {border_color};">
                                        <td><div class="firma-name">{firma}</div></td>
                                        <td><span class="beyanname-badge" style="background: #4caf50;">{kod_0_sayisi} beyanname</span></td>
                                        <td><span class="beyanname-badge" style="background: #2196f3;">{kod_6_sayisi} beyanname</span></td>
                                        <td><span class="beyanname-badge" style="background: {border_color};">KOD {hatali_kod}</span></td>
                                        <td><span class="beyanname-badge" style="background: {border_color};">{hatali_beyanname_sayisi} hatalı</span></td>
                                    </tr>
                        """
                    
                    html += """
                                </tbody>
                            </table>
                    """
                    
                    # Tutarsız beyanname detayları
                    html += """
                            <h3 style="color: #2c3e50; margin: 20px 0;">📋 Hatalı Beyanname Detayları</h3>
                            <div style="background: #fff3cd; border: 1px solid #ffeaa7; border-radius: 6px; padding: 15px; margin: 15px 0;">
                                <strong>⚠️ Dikkat:</strong> Bu tabloda her beyanname <strong>sadece bir kez</strong> gösterilir. 
                                Beyanname birden fazla kalem içeriyorsa kalem sayısı belirtilir.
                            </div>
                            <table class="detail-table">
                                <thead>
                                    <tr>
                                        <th style="width: 30%;">Gönderici Firma</th>
                                        <th style="width: 25%;">Beyanname No<br><small>(Benzersiz)</small></th>
                                        <th style="width: 15%;">Kalem Sayısı</th>
                                        <th style="width: 15%;">Kullandığı İlişki Kodu</th>
                                        <th style="width: 15%;">Doğru Kod</th>
                                    </tr>
                                </thead>
                                <tbody>
                    """
                    
                    # Benzersiz beyanname listesi oluştur
                    unique_beyannames_data = []
                    
                    # Her beyanname için kalem sayısını hesapla
                    for beyanname_no in data_df['Beyanname_no'].unique()[:50]:  # İlk 50 benzersiz beyanname
                        beyanname_rows = data_df[data_df['Beyanname_no'] == beyanname_no]
                        first_row = beyanname_rows.iloc[0]  # İlk satırı al
                        kalem_sayisi = len(beyanname_rows)  # Kalem sayısı
                        
                        unique_beyannames_data.append({
                            'beyanname_no': beyanname_no,
                            'firma': first_row.get(sender_column, ''),
                            'kullandigi_kod': first_row.get('Alici_satici_iliskisi', ''),
                            'dogru_kod': first_row.get('Dogru_Kod', ''),
                            'kalem_sayisi': kalem_sayisi
                        })
                    
                    # Benzersiz beyannameleri göster
                    for beyanname_data in unique_beyannames_data:
                        firma = beyanname_data['firma']
                        beyanname_no = beyanname_data['beyanname_no']
                        kullandigi_kod = beyanname_data['kullandigi_kod']
                        dogru_kod = beyanname_data['dogru_kod']
                        kalem_sayisi = beyanname_data['kalem_sayisi']
                        
                        # Bu firmaya ait istatistikleri bul
                        firma_stats = stats_df[stats_df['Firma'] == firma]
                        if not firma_stats.empty:
                            kod_0_sayisi = firma_stats.iloc[0]['Kod_0_Sayısı']
                            kod_6_sayisi = firma_stats.iloc[0]['Kod_6_Sayısı']
                            firma_info = f"(Toplam: {kod_0_sayisi} benzersiz beyanname kod-0, {kod_6_sayisi} benzersiz beyanname kod-6)"
                        else:
                            firma_info = ""
                        
                        # Kalem sayısı badge rengi
                        if kalem_sayisi == 1:
                            kalem_color = "#6c757d"  # Gri - tek kalem
                        elif kalem_sayisi <= 5:
                            kalem_color = "#17a2b8"  # Mavi - az kalem
                        elif kalem_sayisi <= 10:
                            kalem_color = "#ffc107"  # Sarı - orta kalem
                        else:
                            kalem_color = "#dc3545"  # Kırmızı - çok kalem
                        
                        html += f"""
                                    <tr>
                                        <td>
                                            <div class="firma-name">{firma}</div>
                                            <small style="color: #7f8c8d;">{firma_info}</small>
                                        </td>
                                        <td><span class="beyanname-badge">{beyanname_no}</span></td>
                                        <td><span class="beyanname-badge" style="background: {kalem_color};">{kalem_sayisi} kalem</span></td>
                                        <td><span class="beyanname-badge" style="background: #e74c3c;">KOD {kullandigi_kod}</span></td>
                                        <td><span class="beyanname-badge" style="background: #27ae60;">KOD {dogru_kod}</span></td>
                                    </tr>
                        """
                    
                    total_unique_beyannames = len(data_df['Beyanname_no'].unique())
                    if total_unique_beyannames > 50:
                        html += f"""
                                    <tr>
                                        <td colspan="5" style="text-align: center; font-style: italic; padding: 15px; background-color: #f5f5f5;">
                                            <strong>📋 Toplam {total_unique_beyannames} benzersiz hatalı beyanname bulundu, 
                                            ilk 50 tanesi gösterilmektedir.</strong>
                                        </td>
                                    </tr>
                        """
                    
                    html += """
                                </tbody>
                            </table>
                    """
            
            # Özet tablosu varsa ekle
            if not summary_df.empty:
                html += f"""
                        <h3 style="color: #2c3e50; margin: 20px 0;">📊 Özet Tablo</h3>
                        {summary_df.to_html(index=False, classes="summary-table", escape=False)}
                """
            
            # Eylem gerektiren durumlar
            if problematic_beyannames > 0:
                html += f"""
                        <div class="action-required">
                            <h3>⚠️ Eylem Gerektiren Durumlar</h3>
                            <ul>
                                <li><strong>{problematic_beyannames:,} beyanname</strong> sorunlu olarak tespit edildi</li>
                                <li>Bu beyannameler detaylı incelenmeli</li>
                                <li>İlişki durumu beyanları gözden geçirilmeli</li>
                                <li>Gerekirse düzeltme işlemleri yapılmalı</li>
                            </ul>
                </div>
                """
            
            # Öneriler
            html += f"""
                        <div class="recommendation">
                            <h3>📋 Yönetici Önerileri</h3>
            """
            
            if problematic_beyannames > 100:
                html += """
                            <p><strong>Acil Eylem Gerekli!</strong></p>
                            <ul>
                                <li>Sorunlu beyannamelerin tüm işlemlerini gözden geçirin</li>
                                <li>İlişki durumu beyan süreçlerini iyileştirin</li>
                                <li>Personel eğitimi düzenleyin</li>
                            </ul>
                """
            elif problematic_beyannames > 0:
                html += """
                            <p><strong>Dikkatli İzleme Gerekli.</strong></p>
                            <ul>
                                <li>Sorunlu beyannameleri yakından takip edin</li>
                                <li>Düzenli kontroller yapın</li>
                                <li>Önleyici tedbirler alın</li>
                            </ul>
                """
            else:
                html += """
                            <p><strong>Mükemmel Performans!</strong></p>
                            <ul>
                                <li>Hiçbir sorun tespit edilmedi</li>
                                <li>Mevcut kontrol sistemini sürdürün</li>
                                <li>Düzenli izleme yapın</li>
                            </ul>
                """
            
            html += """
                        </div>
                    </div>
                </div>
            </body>
            </html>
            """
            
            return html
            
        except Exception as e:
            print(f"HTML rapor oluşturma detay hatası: {str(e)}")
            # Basit fallback HTML
            return f"""
            <!DOCTYPE html>
            <html>
            <head><meta charset="UTF-8"><title>Alıcı-Satıcı İlişki Kontrolü</title></head>
            <body>
                <h1>Alıcı-Satıcı İlişki Kontrolü</h1>
                <p>Rapor: {result.get('message', 'Sonuç bilinmiyor')}</p>
                <p>Tip: {result.get('type', 'Bilinmiyor')}</p>
                <p>HTML rapor oluşturulurken hata oluştu: {str(e)}</p>
            </body>
            </html>
            """

    def show_merged_dataframes(self):
        """Tüm dataframe'leri birleştir ve göster"""
        if not self.all_dataframes:
            QMessageBox.warning(self, "Uyarı", "Birleştirilecek veri bulunamadı")
            return
            
        try:
            # İlerleme çubuğunu göster
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            self.status_label.setText("Veriler birleştiriliyor...")
            QApplication.processEvents()
            
            # Tüm dataframe'leri birleştir
            dfs_to_merge = list(self.all_dataframes.values())
            
            if len(dfs_to_merge) == 1:
                # Tek bir dataframe varsa direkt onu göster
                merged_df = dfs_to_merge[0].copy()
            else:
                # Birden fazla dataframe'i birleştir
                merged_df = pd.concat(dfs_to_merge, ignore_index=True)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(50)
            self.status_label.setText("Birleştirilmiş veri gösteriliyor...")
            QApplication.processEvents()
            
            # Birleştirilmiş veriyi göster
            self.merged_df = merged_df
            self.display_dataframe(merged_df)
            
            # Başlık bilgisini değiştir
            file_count = len(self.all_dataframes)
            row_count = len(merged_df)
            self.setWindowTitle(f"Beyanname Kontrol Uygulaması - {file_count} Dosya Birleştirildi ({row_count} Satır)")
            
            # Dosya seçiciyi devre dışı bırak (birleştirme modunda)
            self.file_selector.setEnabled(False)
            
            # İlerleme çubuğunu gizle
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"Toplam {file_count} dosya birleştirildi. Toplam {row_count} satır, {len(merged_df.columns)} sütun.")
            
            # Dashboard'u da güncelle (birleştirilmiş veriyle)
            self.update_dashboard()
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Veriler birleştirilirken hata oluştu: {str(e)}")
            self.progress_bar.setVisible(False)
            self.status_label.setText("Hata oluştu")
            return

    def create_gtip_summary(self):
        """Create GTIP summary table"""
        QMessageBox.information(self, "Bilgi", "Bu özellik henüz tamamlanmamıştır")
    
    def create_country_summary(self):
        """Create country summary pivot table"""
        QMessageBox.information(self, "Bilgi", "Bu özellik henüz tamamlanmamıştır")
    
    def create_rejim_summary(self):
        """Create rejim summary pivot table"""
        QMessageBox.information(self, "Bilgi", "Bu özellik henüz tamamlanmamıştır")
    
    def create_gtip_country_cross(self):
        """Create GTIP-country cross pivot table"""
        QMessageBox.information(self, "Bilgi", "Bu özellik henüz tamamlanmamıştır")

    # Örnekleme fonksiyonları
    def start_sampling(self):
        """Örnekleme işlemini başlatır"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Örnekleme yapılacak veri yüklenmemiş")
            return
        
        # Onay al
        reply = QMessageBox.question(
            self, 
            "Örnekleme Başlat", 
            "Tüm örnekleme kriterleri uygulanarak beyanname seçimi yapılacaktır. Devam etmek istiyor musunuz?",
            QMessageBox.Yes | QMessageBox.No, 
            QMessageBox.Yes
        )
        
        if reply == QMessageBox.No:
            return
        
        try:
            # İşleme başladığını göster
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            self.status_label.setText("Örnekleme yapılıyor...")
            QApplication.processEvents()
            
            # Parametreleri al
            sample_rate = float(self.sample_rate_combo.currentText()) / 100
            min_count = self.min_sample_spin.value()
            max_count = self.max_sample_spin.value()
            
            # İlerleme
            self.progress_bar.setValue(30)
            self.status_label.setText("Beyannameler seçiliyor...")
            QApplication.processEvents()
            
            # Örnekleme yap
            self.sampling_tool.set_dataframe(self.current_df)
            
            # İlerleme
            self.progress_bar.setValue(50)
            QApplication.processEvents()
            
            results_df = self.sampling_tool.run_sampling(
                min_sample_count=min_count,
                max_sample_count=max_count,
                sample_percentage=sample_rate
            )
            
            # İlerleme
            self.progress_bar.setValue(80)
            self.status_label.setText("Sonuçlar gösteriliyor...")
            QApplication.processEvents()
            
            # Sonuçları göster
            self.sampling_viewer.set_dataframe(results_df)
            
            # İstatistikleri güncelle
            total_beyannames = self.sampling_tool.sampling_stats.get('total_beyannames', 0)
            selected_count = len(self.sampling_tool.selected_beyannames)
            target_count = self.sampling_tool.sampling_stats.get('target_sample_count', 0)
            selection_rate = (selected_count / total_beyannames * 100) if total_beyannames > 0 else 0
            
            self.sampling_stats_label.setText(
                f"Toplam {total_beyannames} beyannameden {selected_count} tanesi seçildi. " +
                f"Hedef: {target_count} beyanname. " +
                f"Seçim oranı: %{selection_rate:.2f}"
            )
            
            # Butonları etkinleştir
            self.export_excel_btn.setEnabled(True)
            self.clear_sampling_btn.setEnabled(True)
            
            # İşlemi tamamla
            self.progress_bar.setValue(100)
            self.status_label.setText(f"Örnekleme tamamlandı. {selected_count} beyanname seçildi.")
            self.progress_bar.setVisible(False)
            
            # Örnekleme sekmesine geçiş yap
            self.tabs.setCurrentWidget(self.sampling_tab)
            
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            QMessageBox.critical(self, "Hata", f"Örnekleme sırasında hata oluştu: {str(e)}")
            self.status_label.setText("Örnekleme sırasında hata oluştu")
            self.progress_bar.setVisible(False)
    
    def export_sampling_to_excel(self):
        """Örnekleme sonuçlarını Excel dosyasına aktarır - Arka plan thread ile"""
        if not hasattr(self, 'sampling_tool') or not self.sampling_tool.selected_beyannames:
            QMessageBox.warning(self, "Uyarı", "Dışa aktarılacak örnekleme sonucu bulunamadı")
            return
        
        try:
            # Excel dosyası adı için kullanıcıdan dosya yolu al
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Excel Dosyasını Kaydet", "", "Excel Files (*.xlsx)"
            )
            
            if not file_path:
                return
            
            # Uzantı kontrolü
            if not file_path.endswith('.xlsx'):
                file_path += '.xlsx'
            
            # Dosya kullanımda mı kontrol et
            try:
                with open(file_path, 'a') as test_file:
                    pass
            except PermissionError:
                QMessageBox.critical(self, "Hata", 
                    f"Dosya başka bir program tarafından kullanılıyor veya açık durumda: {file_path}\n"
                    "Lütfen dosyayı kapatıp tekrar deneyin."
                )
                self.progress_bar.setVisible(False)
                self.status_label.setText("Excel'e aktarma başarısız: Dosya açık")
                return
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"Dosya izin hatası: {str(e)}")
                self.progress_bar.setVisible(False)
                self.status_label.setText("Excel'e aktarma başarısız: Dosya hatası")
                return
            
            # Örneklemeyi kontrol et
            if len(self.sampling_tool.selected_beyannames) == 0:
                QMessageBox.warning(self, "Uyarı", "Seçilen beyanname bulunmuyor.")
                self.progress_bar.setVisible(False)
                self.status_label.setText("Excel'e aktarma başarısız: Seçilen beyanname yok")
                return
            
            # İlerleme çubuğunu göster
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            self.status_label.setText("Excel aktarma işlemi başlatılıyor...")
            QApplication.processEvents()
            
            # İptal butonu göster (QDialog içinde)
            self.cancel_dialog = QDialog(self)
            self.cancel_dialog.setWindowTitle("Excel'e Aktarma")
            self.cancel_dialog.setModal(True)
            self.cancel_dialog.setFixedSize(300, 100)
            
            dialog_layout = QVBoxLayout(self.cancel_dialog)
            dialog_layout.addWidget(QLabel("Excel aktarım işlemi çalışıyor...\nBu işlem veri boyutuna göre zaman alabilir."))
            
            cancel_btn = QPushButton("İptal")
            cancel_btn.clicked.connect(self.cancel_excel_export)
            dialog_layout.addWidget(cancel_btn)
            
            # Thread oluştur ve başlat
            self.excel_thread = ExcelExportThread(self.sampling_tool, file_path)
            
            # Thread sinyallerini bağla
            self.excel_thread.progress.connect(self.update_excel_progress)
            self.excel_thread.finished.connect(self.on_excel_export_finished)
            
            # Thread'i başlat
            self.excel_thread.start()
            
            # Dialog'u göster (non-blocking)
            self.cancel_dialog.show()
            
        except Exception as e:
            # Beklenmeyen hata durumunda
            import traceback
            print(f"Excel'e aktarma sırasında beklenmeyen hata: {str(e)}")
            print(traceback.format_exc())
            
            QMessageBox.critical(self, "Hata", f"Excel'e aktarma sırasında beklenmeyen bir hata oluştu: {str(e)}")
            self.status_label.setText("Excel'e aktarma sırasında beklenmeyen hata oluştu")
            self.progress_bar.setVisible(False)

    def update_excel_progress(self, value, message):
        """Excel export progress güncellemesi"""
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setValue(value)
        if hasattr(self, 'status_label'):
            self.status_label.setText(message)
        QApplication.processEvents()

    def on_excel_export_finished(self, success, message, file_path):
        """Excel export thread tamamlandığında"""
        # İlerleme çubuğunu güncelle
        if hasattr(self, 'progress_bar'):
            self.progress_bar.setValue(100 if success else 0)
            self.progress_bar.setVisible(False)
        
        # Durum mesajını güncelle
        if hasattr(self, 'status_label'):
            self.status_label.setText(message)
        
        # İptal dialog'unu kapat
        if hasattr(self, 'cancel_dialog') and self.cancel_dialog:
            self.cancel_dialog.accept()
            self.cancel_dialog = None
        
        # Sonuç mesajını göster
        if success:
            QMessageBox.information(self, "İşlem Tamamlandı", f"Örnekleme sonuçları başarıyla Excel dosyasına aktarıldı:\n{file_path}")
        else:
            QMessageBox.warning(self, "İşlem Başarısız", message)

    def cancel_excel_export(self):
        """Excel export işlemini iptal et"""
        if hasattr(self, 'excel_thread') and self.excel_thread.isRunning():
            reply = QMessageBox.question(self, "İşlemi İptal Et", 
                                        "Excel aktarma işlemini iptal etmek istediğinizden emin misiniz?",
                                        QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                # Thread'i iptal et
                self.excel_thread.cancel()
                
                # Dialog'u kapat
                if hasattr(self, 'cancel_dialog') and self.cancel_dialog:
                    self.cancel_dialog.accept()
                    self.cancel_dialog = None
                
                # İlerleme çubuğunu güncelle
                if hasattr(self, 'progress_bar'):
                    self.progress_bar.setVisible(False)
                
                # Durum mesajını güncelle
                if hasattr(self, 'status_label'):
                    self.status_label.setText("Excel aktarma işlemi iptal edildi")
                
                QMessageBox.information(self, "İşlem İptal Edildi", "Excel aktarma işlemi kullanıcı tarafından iptal edildi.")
        
    def clear_sampling(self):
        """Örnekleme sonuçlarını temizler"""
        if hasattr(self, 'sampling_tool'):
            # Örnekleme aracını temizle
            self.sampling_tool.selected_beyannames = set()
            self.sampling_tool.selection_reasons = {}
            
            # UI'ı temizle
            self.sampling_viewer.set_dataframe(None)
            self.sampling_stats_label.setText("Örnekleme henüz yapılmadı")
            
            # Butonları devre dışı bırak
            self.export_excel_btn.setEnabled(False)
            self.clear_sampling_btn.setEnabled(False)
            
            # Bilgi ver
            self.status_label.setText("Örnekleme sonuçları temizlendi")

    def setup_application_stability(self):
        """Uygulama stabilitesini arttırıcı ayarlar yapar"""
        # Ana uygulama için 5 dakikada bir otomatik kurtarma
        self.stability_timer = QTimer(self)
        self.stability_timer.timeout.connect(self.check_application_health)
        self.stability_timer.start(300000)  # 5 dakika
        
        # Hata yakalama hook'u
        sys.excepthook = self.global_exception_handler

    def check_application_health(self):
        """Uygulama sağlık kontrolü"""
        try:
            # Bellek ve diğer kaynakları temizle
            import gc
            gc.collect()
            
            # UI ile ilgili tüm bekleyen olayları işle
            QApplication.processEvents()
        except Exception as e:
            print(f"Sağlık kontrolü hatası: {str(e)}")

    def global_exception_handler(self, exctype, value, traceback):
        """Global exception handler - uygulama kapanmasını önler"""
        try:
            # Hatayı logla
            import traceback as tb
            error_msg = ''.join(tb.format_exception(exctype, value, traceback))
            print(f"Yakalanan hata:\n{error_msg}")
            
            # Kullanıcıya hatayı bildir ama kapanmayı önle
            QMessageBox.critical(
                self, 
                "Uygulama Hatası", 
                "Bir hata oluştu, ancak uygulama çalışmaya devam edecek.\n\n"
                f"Hata detayı: {str(value)}"
            )
        except:
            # Çift hata durumundan kaçın
            pass

    def check_gtip_urun_kodu_consistency(self):
        """GTİP-Ürün Kodu tutarlılık kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("GTİP-Ürün Kodu tutarlılık kontrolü çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()  # UI güncellemesi için
            
            # Analiz öncesi veri büyüklüğünü kontrol et
            if len(self.current_df) > 1000:
                # Büyük veri setleri için örnekleme yap
                sample_size = min(1000, len(self.current_df))
                sample_df = self.current_df.sample(sample_size)
                process_df = sample_df
                self.status_label.setText(f"Büyük veri seti ({len(self.current_df)} satır): Analiz için {sample_size} satır örnekleniyor...")
            else:
                process_df = self.current_df
            
            QApplication.processEvents()  # UI güncellemesi için
            
            # Analiz fonksiyonunu çağır
            gtip_urun_kodu_check = check_gtip_urun_kodu_consistency(process_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(50)
            QApplication.processEvents()
            
            if gtip_urun_kodu_check is not None:
                result = {
                    "GTIP-Ürün Kodu tutarlılık kontrolü": gtip_urun_kodu_check
                }
                
                # Sonuçları widget'a aktar
                self.check_results_widget.set_check_results(result)
                
                # İlerleme çubuğunu güncelle
                self.progress_bar.setValue(80)
                QApplication.processEvents()
                
                # HTML raporunu CheckResultsWidget'a ekle (arka planda)
                if "html_report" in gtip_urun_kodu_check:
                    print("HTML raporu bulundu, gösteriliyor...")
                    
                    # Görsel rapor için ekran boyutunu optimize et
                    try:
                        # Panelin genişliğini mümkün olduğunca artır
                        self.check_results_widget.html_view.setSizePolicy(
                            QSizePolicy.Expanding, QSizePolicy.Expanding)
                    except Exception as e:
                        print(f"Görünüm boyutu ayarlanamadı: {str(e)}")
                    
                    # Önce kısa bir yükleme bildirimi göster
                    loading_html = """
                        <html>
                        <body style="font-family: Arial; padding: 20px; text-align: center;">
                            <h2>GTİP-Ürün Kodu Raporu Hazırlanıyor</h2>
                            <p>Lütfen bekleyin, rapor hazırlanıyor...</p>
                            <div style="width:50%;height:20px;background-color:#f0f0f0;margin:20px auto;border-radius:10px;">
                                <div style="width:100%;height:20px;background-color:#4299e1;border-radius:10px;animation:loading 1.5s infinite;"></div>
                            </div>
                            <style>
                                @keyframes loading {
                                    0% { width: 0%; }
                                    50% { width: 100%; }
                                    100% { width: 0%; }
                                }
                            </style>
                        </body>
                        </html>
                    """
                    self.check_results_widget.html_view.setHtml(loading_html)
                    
                    # Görsel rapor sekmesine doğrudan geçiş yap
                    self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                    QApplication.processEvents()  # UI güncellemesi için
                    
                    # HTML içeriğini gecikmeli yükle (zamanlayıcı ile)
                    from PyQt5.QtCore import QTimer
                    QTimer.singleShot(300, lambda: self.check_results_widget.set_html_report(gtip_urun_kodu_check["html_report"]))
                        
                # Durum bilgisini güncelle
                status = gtip_urun_kodu_check.get("status", "")
                if status == "warning":
                    self.status_label.setText("GTİP-Ürün Kodu tutarsızlıkları bulundu")
                else:
                    self.status_label.setText("GTİP-Ürün Kodu kontrolü tamamlandı")
            else:
                result = {
                    "GTIP-Ürün Kodu tutarlılık kontrolü": {
                        "status": "ok",
                        "message": "GTIP-Ürün Kodu tutarlılık kontrolü bulunamadı."
                    }
                }
                self.check_results_widget.set_check_results(result)
                self.status_label.setText("GTİP-Ürün Kodu kontrolü tamamlandı")
                
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"GTİP-Ürün Kodu kontrolü sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_rarely_used_currency(self):
        """Nadiren kullanılan döviz kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Nadiren kullanılan döviz analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Analiz fonksiyonunu çağır
            currency_check = check_rarely_used_currency(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if currency_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", currency_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(currency_check["message"])
                return
            
            result = {
                "Nadiren Kullanılan Döviz Analizi": currency_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in currency_check:
                self.check_results_widget.set_html_report(currency_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in currency_check and "summary" in currency_check:
                self.check_results_widget.show_details(currency_check["data"])
                self.check_results_widget.show_summary(currency_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(currency_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Nadiren kullanılan döviz analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)
            
    def check_rarely_used_origin_country(self):
        """Nadiren kullanılan menşe ülke kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Nadiren kullanılan menşe ülke analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Analiz fonksiyonunu çağır
            country_check = check_rarely_used_origin_country(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if country_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", country_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(country_check["message"])
                return
            
            result = {
                "Nadiren Kullanılan Menşe Ülke Analizi": country_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in country_check:
                self.check_results_widget.set_html_report(country_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in country_check and "summary" in country_check:
                self.check_results_widget.show_details(country_check["data"])
                self.check_results_widget.show_summary(country_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(country_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Nadiren kullanılan menşe ülke analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)
            
    def check_rarely_used_payment_method(self):
        """Nadiren kullanılan ödeme şekli kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Nadiren kullanılan ödeme şekli analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Analiz fonksiyonunu çağır
            payment_check = check_rarely_used_payment_method(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if payment_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", payment_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(payment_check["message"])
                return
            
            result = {
                "Nadiren Kullanılan Ödeme Şekli Analizi": payment_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in payment_check:
                self.check_results_widget.set_html_report(payment_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in payment_check and "summary" in payment_check:
                self.check_results_widget.show_details(payment_check["data"])
                self.check_results_widget.show_summary(payment_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(payment_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Nadiren kullanılan ödeme şekli analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_unit_price_increase(self):
        """Birim fiyat artış kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Birim fiyat artış analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Analiz fonksiyonunu çağır
            price_check = check_unit_price_increase(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if price_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", price_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(price_check["message"])
                return
            
            result = {
                "Birim Fiyat Artış Analizi": price_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in price_check:
                self.check_results_widget.set_html_report(price_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in price_check and "summary" in price_check:
                self.check_results_widget.show_details(price_check["data"])
                self.check_results_widget.show_summary(price_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(price_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            import traceback
            error_details = traceback.format_exc()
            error_msg = f"Birim fiyat artış analizi sırasında hata: {str(e)}"
            print(error_msg)
            print(f"Hata detayı: {error_details}")
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_kdv_consistency(self):
        """KDV Kontrol - Aynı GTİP kodunda farklı Vergi_2_Oran değerleri kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("KDV kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Yeni KDV kontrol fonksiyonunu çağır
            from analysis_modules.kdv_kontrol import check_kdv_kontrol
            kdv_check = check_kdv_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if kdv_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", kdv_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(kdv_check["message"])
                return
            
            result = {
                "KDV Kontrol Analizi": kdv_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in kdv_check:
                self.check_results_widget.set_html_report(kdv_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in kdv_check and "summary" in kdv_check:
                self.check_results_widget.show_details(kdv_check["data"])
                self.check_results_widget.show_summary(kdv_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(kdv_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"KDV kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_domestic_expense_variation(self):
        """Yurt İçi Gider Kontrol - Beyanname bazında gider ve ağırlık analizi"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Yurt içi gider kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Yeni yurt içi gider kontrol fonksiyonunu çağır
            from analysis_modules.yurt_ici_gider_kontrol import check_yurt_ici_gider_kontrol
            expense_check = check_yurt_ici_gider_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if expense_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", expense_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(expense_check["message"])
                return
            
            result = {
                "Yurt İçi Gider Kontrol Analizi": expense_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in expense_check:
                self.check_results_widget.set_html_report(expense_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in expense_check and "summary" in expense_check:
                self.check_results_widget.show_details(expense_check["data"])
                self.check_results_widget.show_summary(expense_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(expense_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Yurt içi gider kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_foreign_expense_variation(self):
        """Yurt Dışı Gider Kontrol - İki farklı kontrol analizi"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Yurt dışı gider kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Yeni yurt dışı gider kontrol fonksiyonunu çağır
            from analysis_modules.yurt_disi_gider_kontrol import check_yurt_disi_gider_kontrol
            expense_check = check_yurt_disi_gider_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if expense_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", expense_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(expense_check["message"])
                return
            
            result = {
                "Yurt Dışı Gider Kontrol Analizi": expense_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in expense_check:
                self.check_results_widget.set_html_report(expense_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in expense_check and "summary" in expense_check:
                self.check_results_widget.show_details(expense_check["data"])
                self.check_results_widget.show_summary(expense_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(expense_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Yurt dışı gider kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_supalan_storage(self):
        """Supalan Depolama Kontrol - TAŞIT ÜSTÜ - SUPALAN SAHASI depolama gideri kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Supalan depolama kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Supalan depolama kontrol fonksiyonunu çağır
            supalan_check = check_supalan_depolama_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if supalan_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", supalan_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(supalan_check["message"])
                return
            
            result = {
                "Supalan Depolama Kontrol": supalan_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in supalan_check:
                self.check_results_widget.set_html_report(supalan_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in supalan_check and "summary" in supalan_check:
                self.check_results_widget.show_details(supalan_check["data"])
                self.check_results_widget.show_summary(supalan_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(supalan_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Supalan depolama kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_igv_consistency(self):
        """IGV (İlave Gümrük Vergisi) tutarlılığını kontrol eder"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("IGV kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # IGV kontrol fonksiyonunu çağır
            from analysis_modules.igv_analysis import check_igv_consistency as igv_analysis_func
            result = igv_analysis_func(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if not result["success"]:
                QMessageBox.warning(self, "Uyarı", result["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(result["message"])
                return
            
            # HTML raporu oluştur
            html_report = self._generate_igv_html_report(result)
            
            # Sonuçları widget'a aktar
            check_result = {
                "status": "success",
                "message": result["message"],
                "data": result["data"],
                "summary": result.get("summary", {}),
                "html_report": html_report
            }
            
            results = {"IGV Kontrol Analizi": check_result}
            self.check_results_widget.set_check_results(results)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            self.check_results_widget.set_html_report(html_report)
            self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if not result["data"].empty:
                self.check_results_widget.show_details(result["data"])
                if "summary" in result:
                    self.check_results_widget.show_summary(result["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(result["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"IGV kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)
    
    def _generate_igv_html_report(self, result):
        """IGV analizi için HTML raporu oluşturur"""
        try:
            data = result["data"]
            summary = result.get("summary", {})
            significant_diff = result.get("significant_differences", pd.DataFrame())
            
            html = f"""
            <html>
            <head>
                <meta charset="UTF-8">
                <title>IGV Kontrol Analizi</title>
                <style>
                    body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 20px; background-color: #f8f9fa; }}
                    .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; }}
                    .summary {{ background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
                    .metric {{ display: inline-block; margin: 10px 20px 10px 0; padding: 15px; background: #f8f9fa; border-radius: 8px; border-left: 4px solid #007bff; }}
                    .metric-value {{ font-size: 24px; font-weight: bold; color: #007bff; }}
                    .metric-label {{ font-size: 14px; color: #6c757d; margin-top: 5px; }}
                    .table-container {{ background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
                    table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
                    th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #dee2e6; }}
                    th {{ background-color: #f8f9fa; font-weight: 600; color: #495057; }}
                    tr:hover {{ background-color: #f8f9fa; }}
                    .positive {{ color: #28a745; font-weight: bold; }}
                    .negative {{ color: #dc3545; font-weight: bold; }}
                    .warning {{ color: #ffc107; font-weight: bold; }}
                    .info {{ color: #17a2b8; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>💰 IGV (İlave Gümrük Vergisi) Kontrol Analizi</h1>
                    <p>IGV listesindeki GTIP kodları ile sistemdeki vergi hesaplamalarının karşılaştırılması</p>
                </div>
            """
            
            # Özet istatistikler
            if summary:
                html += f"""
                <div class="summary">
                    <h2>📊 Özet İstatistikler</h2>
                    <div class="metric">
                        <div class="metric-value">{summary.get('toplam_igv_kayit', 0)}</div>
                        <div class="metric-label">Toplam IGV Kaydı</div>
                    </div>
                    <div class="metric">
                        <div class="metric-value">{summary.get('onemli_fark_sayisi', 0)}</div>
                        <div class="metric-label">Önemli Fark (>5%)</div>
                    </div>
                    <div class="metric">
                        <div class="metric-value">{summary.get('ortalama_fark', 0):.2f}</div>
                        <div class="metric-label">Ortalama Fark</div>
                    </div>
                    <div class="metric">
                        <div class="metric-value">{summary.get('toplam_fark', 0):.2f}</div>
                        <div class="metric-label">Toplam Fark</div>
                    </div>
                </div>
                """
            
            # Önemli farklar tablosu
            if not significant_diff.empty:
                html += f"""
                <div class="table-container">
                    <h2>⚠️ Önemli Farklar (>5%)</h2>
                    <table>
                        <thead>
                            <tr>
                                <th>Beyanname No</th>
                                <th>GTIP</th>
                                <th>Menşe Ülke</th>
                                <th>Excel IGV (%)</th>
                                <th>Sistem IGV (%)</th>
                                <th>Beklenen IGV</th>
                                <th>Sistem IGV</th>
                                <th>Fark</th>
                                <th>Fark %</th>
                            </tr>
                        </thead>
                        <tbody>
                """
                
                for _, row in significant_diff.head(20).iterrows():
                    fark_class = "positive" if row.get('Fark', 0) > 0 else "negative"
                    html += f"""
                    <tr>
                        <td>{row.get('Beyanname_no', '')}</td>
                        <td>{row.get('Gtip', '')}</td>
                        <td>{row.get('Mensei_ulke', '')}</td>
                        <td class="info">{row.get('Excel_IGV_orani', 0):.1f}%</td>
                        <td class="warning">{row.get('Sistem_IGV_orani', 0):.1f}%</td>
                        <td>{row.get('Beklenen_IGV_tutari', 0):.2f}</td>
                        <td>{row.get('Sistem_IGV_tutari', 0):.2f}</td>
                        <td class="{fark_class}">{row.get('Fark', 0):.2f}</td>
                        <td class="{fark_class}">{row.get('Fark_yuzdesi', 0):.1f}%</td>
                    </tr>
                    """
                
                html += """
                        </tbody>
                    </table>
                </div>
                """
            
            # Tüm sonuçlar tablosu
            if not data.empty:
                html += f"""
                <div class="table-container">
                    <h2>📋 Tüm IGV Kontrol Sonuçları</h2>
                    <p>Toplam {len(data)} kayıt analiz edildi.</p>
                    <table>
                        <thead>
                            <tr>
                                <th>Beyanname No</th>
                                <th>GTIP</th>
                                <th>Menşe Ülke</th>
                                <th>Vergi Matrahı</th>
                                <th>Excel IGV (%)</th>
                                <th>Sistem IGV (%)</th>
                                <th>Beklenen IGV</th>
                                <th>Sistem IGV</th>
                                <th>Fark</th>
                                <th>Kullanılan Sütun</th>
                            </tr>
                        </thead>
                        <tbody>
                """
                
                for _, row in data.head(50).iterrows():
                    fark_class = ""
                    if pd.notna(row.get('Fark')):
                        if abs(row.get('Fark', 0)) > 10:
                            fark_class = "negative"
                        elif abs(row.get('Fark', 0)) > 5:
                            fark_class = "warning"
                        else:
                            fark_class = "positive"
                    
                    html += f"""
                    <tr>
                        <td>{row.get('Beyanname_no', '')}</td>
                        <td>{row.get('Gtip', '')}</td>
                        <td>{row.get('Mensei_ulke', '')}</td>
                        <td>{row.get('Vergi_matrahi', 0):.2f}</td>
                        <td class="info">{row.get('Excel_IGV_orani', 0):.1f}%</td>
                        <td class="warning">{row.get('Sistem_IGV_orani', 0):.1f}%</td>
                        <td>{row.get('Beklenen_IGV_tutari', 0):.2f}</td>
                        <td>{row.get('Sistem_IGV_tutari', 0):.2f}</td>
                        <td class="{fark_class}">{row.get('Fark', 0):.2f}</td>
                        <td class="info">{row.get('Kullanilan_sutun', '')}</td>
                    </tr>
                    """
                
                html += """
                        </tbody>
                    </table>
                </div>
                """
            
            html += """
            </body>
            </html>
            """
            
            return html
            
        except Exception as e:
            return f"<html><body><h1>Rapor oluşturma hatası: {str(e)}</h1></body></html>"

    def setup_shortcuts(self):
        """Uygulama için kısayol tuşlarını ayarla"""
        from PyQt5.QtGui import QKeySequence
        from PyQt5.QtWidgets import QShortcut
        
        # Alt+F ile arama kısayolu
        search_shortcut = QShortcut(QKeySequence("Alt+F"), self)
        search_shortcut.activated.connect(self.show_search_dialog)
    
    def show_search_dialog(self):
        """Alt+F basıldığında metin arama dialog'unu göster"""
        from PyQt5.QtWidgets import QDialog, QLineEdit, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox
        
        # Aktif görünümdeki tabloyu veya veri görüntüleyiciyi bul
        active_viewer = self.get_active_data_viewer()
        if not active_viewer:
            QMessageBox.information(self, "Bilgi", "Arama yapılacak aktif bir tablo bulunamadı.")
            return
        
        # Arama dialog'u oluştur
        search_dialog = QDialog(self)
        search_dialog.setWindowTitle("Metin Ara")
        search_dialog.setFixedWidth(400)
        search_dialog.setStyleSheet("""
            QLabel {
                font-weight: bold;
            }
            QLineEdit {
                padding: 5px;
                border: 1px solid #ccc;
                border-radius: 3px;
            }
            QPushButton {
                padding: 5px 15px;
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton#closeButton {
                background-color: #f44336;
            }
            QPushButton#closeButton:hover {
                background-color: #d32f2f;
            }
        """)
        
        # Dialog layoutu
        layout = QVBoxLayout(search_dialog)
        
        # Arama metni girişi
        search_layout = QHBoxLayout()
        search_label = QLabel("Aranacak metin:")
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Aramak istediğiniz metni girin...")
        self.search_input.returnPressed.connect(lambda: self.search_in_table(active_viewer))
        search_layout.addWidget(search_label)
        search_layout.addWidget(self.search_input)
        layout.addLayout(search_layout)
        
        # Büyük/küçük harf ve tam kelime eşleşmesi seçenekleri
        options_layout = QHBoxLayout()
        self.case_sensitive = QCheckBox("Büyük/küçük harf duyarlı")
        self.whole_word = QCheckBox("Tam kelime eşleşmesi")
        options_layout.addWidget(self.case_sensitive)
        options_layout.addWidget(self.whole_word)
        layout.addLayout(options_layout)
        
        # Arama sonuç bilgisi
        self.result_label = QLabel("Arama yapmak için metni girin.")
        layout.addWidget(self.result_label)
        
        # Düğmeler
        buttons_layout = QHBoxLayout()
        find_button = QPushButton("Bul")
        find_button.clicked.connect(lambda: self.search_in_table(active_viewer))
        
        find_next_button = QPushButton("Sonraki")
        find_next_button.clicked.connect(lambda: self.find_next(active_viewer))
        
        close_button = QPushButton("Kapat")
        close_button.setObjectName("closeButton")
        close_button.clicked.connect(search_dialog.close)
        
        buttons_layout.addWidget(find_button)
        buttons_layout.addWidget(find_next_button)
        buttons_layout.addStretch()
        buttons_layout.addWidget(close_button)
        layout.addLayout(buttons_layout)
        
        # Arama sonuçları için değişkenleri tanımla
        self.search_results = []
        self.current_result_index = -1
        
        # Dialog'u göster
        search_dialog.exec_()
    
    def get_active_data_viewer(self):
        """Aktif durumda olan veri görüntüleyici widget'ı bul"""
        # Ana veri görüntüleyici kontrolü
        if hasattr(self, 'data_viewer') and self.data_viewer.isVisible():
            return self.data_viewer
            
        # CheckResultsWidget'taki tablo kontrolü
        if hasattr(self, 'check_results_widget') and self.check_results_widget.isVisible():
            if hasattr(self.check_results_widget, 'details_view') and self.check_results_widget.details_view.isVisible():
                return self.check_results_widget.details_view
        
        # Örnekleme görüntüleyici kontrolü
        if hasattr(self, 'sampling_viewer') and self.sampling_viewer.isVisible():
            return self.sampling_viewer
            
        return None
    
    def search_in_table(self, viewer):
        """Tabloda metin ara"""
        if not viewer or not hasattr(viewer, 'table_view'):
            return
            
        search_text = self.search_input.text()
        if not search_text:
            self.result_label.setText("Aranacak metin girilmedi.")
            return
            
        # Modeli al
        model = viewer.table_view.model()
        if not model:
            self.result_label.setText("Tablo modeli bulunamadı.")
            return
            
        # Arama parametreleri
        case_sensitive = self.case_sensitive.isChecked()
        whole_word = self.whole_word.isChecked()
        
        # Arama sonuçlarını temizle
        self.search_results = []
        self.current_result_index = -1
        
        # Tabloda ara
        for row in range(model.rowCount()):
            for col in range(model.columnCount()):
                index = model.index(row, col)
                cell_text = str(model.data(index, Qt.DisplayRole) or "")
                
                # Arama kriterlerine göre kontrol et
                found = False
                if case_sensitive:
                    if whole_word:
                        # Tam kelime eşleşmesi, büyük/küçük harf duyarlı
                        words = cell_text.split()
                        found = search_text in words
                    else:
                        # Parça eşleşmesi, büyük/küçük harf duyarlı
                        found = search_text in cell_text
                else:
                    if whole_word:
                        # Tam kelime eşleşmesi, büyük/küçük harf duyarsız
                        words = cell_text.lower().split()
                        found = search_text.lower() in words
                    else:
                        # Parça eşleşmesi, büyük/küçük harf duyarsız
                        found = search_text.lower() in cell_text.lower()
                
                if found:
                    self.search_results.append((row, col))
        
        # Sonuçları göster
        if self.search_results:
            self.result_label.setText(f"{len(self.search_results)} eşleşme bulundu.")
            self.find_next(viewer)  # İlk sonuca git
        else:
            self.result_label.setText("Eşleşme bulunamadı.")
    
    def find_next(self, viewer):
        """Sonraki arama sonucuna git"""
        if not self.search_results:
            return
            
        # Sonraki sonuç indeksini ayarla
        if self.current_result_index < len(self.search_results) - 1:
            self.current_result_index += 1
        else:
            self.current_result_index = 0  # Başa dön
            
        # Sonuca git
        row, col = self.search_results[self.current_result_index]
        self.result_label.setText(f"Eşleşme {self.current_result_index + 1}/{len(self.search_results)}")
        
        # Tabloda hücreyi seç ve görünür yap
        index = viewer.table_view.model().index(row, col)
        viewer.table_view.setCurrentIndex(index)
        viewer.table_view.scrollTo(index)

    def on_main_tab_changed(self, index):
        """Ana tab değişiminde çağrılan fonksiyon"""
        try:
            tab_text = self.tabs.tabText(index)
            
            # Dashboard sekmesine geçişte dashboard'u güncelle
            if tab_text == "Dashboard":
                self.update_dashboard()
            
            # Sadece Veri Görünümü sekmesinde arama paneli açık olsun
            if hasattr(self, 'data_viewer') and hasattr(self.data_viewer, 'filter_group'):
                if tab_text == "Veri Görünümü":
                    self.data_viewer.filter_group.setVisible(True)
                else:
                    self.data_viewer.filter_group.setVisible(False)
        except Exception as e:
            print(f"Tab değişimi hatası: {e}")

    def on_check_results_tab_changed(self, index):
        tab_text = self.check_results_widget.tabs.tabText(index)
        # Sadece Veri Görünümü ve Detaylar sekmesinde arama paneli açık olsun
        if tab_text in ["Veri Görünümü", "Detaylar"]:
            self.check_results_widget.data_tab.layout().itemAt(0).widget().setVisible(True)
        else:
            self.check_results_widget.data_tab.layout().itemAt(0).widget().setVisible(False)

    def on_tab_changed(self, index):
        """Ana tab değişiminde çağrılan fonksiyon"""
        self.on_main_tab_changed(index)

    def check_gtip_tanim_detail(self):
        """GTİP-Tanım Detay Analizi: Aynı ticari tanımda farklı GTIP kodları kullanılan eşyaları detaylı göster"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # Gerekli sütunları kontrol et
            required_columns = ['Gtip', 'Ticari_tanimi']
            missing_columns = [col for col in required_columns if col not in self.current_df.columns]
            
            if missing_columns:
                QMessageBox.warning(self, "Uyarı", f"Gerekli sütunlar eksik: {', '.join(missing_columns)}")
                return
            
            # İşlem başladığını göster
            self.status_label.setText("GTİP-Tanım Detay Analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(20)
            QApplication.processEvents()
            
            # Boş ticari tanımları filtrele
            filtered_df = self.current_df[
                self.current_df['Ticari_tanimi'].notna() & 
                (self.current_df['Ticari_tanimi'] != '') &
                self.current_df['Gtip'].notna() & 
                (self.current_df['Gtip'] != '')
            ].copy()
            
            self.progress_bar.setValue(40)
            QApplication.processEvents()
            
            if len(filtered_df) == 0:
                QMessageBox.information(self, "Bilgi", "Analiz için uygun veri bulunamadı.")
                self.progress_bar.setVisible(False)
                return
            
            # Her ticari tanım için benzersiz GTİP kodlarını bul
            grouped = filtered_df.groupby('Ticari_tanimi')['Gtip'].unique().reset_index()
            grouped['GTİP_Sayısı'] = grouped['Gtip'].apply(len)
            
            # Birden fazla GTİP kodu olan ticari tanımları filtrele
            multiple_gtips = grouped[grouped['GTİP_Sayısı'] > 1].sort_values(by='GTİP_Sayısı', ascending=False)
            
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if multiple_gtips.empty:
                QMessageBox.information(self, "Sonuç", "Aynı ticari tanımda farklı GTİP kodu kullanımı tespit edilmedi.")
                self.progress_bar.setVisible(False)
                return
            
            # Detaylı sonuç DataFrame'i oluştur
            result_rows = []
            
            for _, row in multiple_gtips.iterrows():
                ticari_tanim = row['Ticari_tanimi']
                
                # Bu ticari tanıma sahip tüm kayıtları al
                ticari_rows = filtered_df[filtered_df['Ticari_tanimi'] == ticari_tanim]
                
                for _, data_row in ticari_rows.iterrows():
                    result_row = {
                        'Ticari_tanimi': ticari_tanim,
                        'Gtip': data_row.get('Gtip', ''),
                        'Adi_unvani': data_row.get('Adi_unvani', ''),
                        'Beyanname_no': data_row.get('Beyanname_no', ''),
                    }
                    
                    # Beyanname kalem numarası için farklı sütun isimlerini dene
                    kalem_no = ''
                    for col_name in ['BeyannameKalemNo', 'Kalem_sira_no', 'Kalem_No', 'KalemNo']:
                        if col_name in data_row and pd.notna(data_row[col_name]):
                            kalem_no = str(data_row[col_name])
                            break
                    result_row['BeyannameKalemNo'] = kalem_no
                    
                    # Diğer yararlı sütunları da ekle
                    for col in ['Mensei_ulke', 'Fatura_miktari', 'Fatura_miktarinin_dovizi']:
                        if col in data_row:
                            result_row[col] = data_row[col]
                    
                    result_rows.append(result_row)
            
            result_df = pd.DataFrame(result_rows)
            
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporu oluştur
            html_content = self._generate_gtip_tanim_detail_html(result_df, multiple_gtips)
            
            # Sonuçları göster
            check_result = {
                "GTİP-Tanım Detay Analizi": {
                    "status": "warning" if not result_df.empty else "ok",
                    "message": f"{len(multiple_gtips)} ticari tanımda farklı GTİP kodları tespit edildi. Toplam {len(result_df)} kayıt bulundu.",
                    "data": result_df,
                    "html_report": html_content
                }
            }
            
            self.check_results_widget.set_check_results(check_result)
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"GTİP-Tanım Detay Analizi tamamlandı. {len(multiple_gtips)} ticari tanımda tutarsızlık bulundu.")
            
        except Exception as e:
            error_msg = f"GTİP-Tanım Detay Analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def _generate_gtip_tanim_detail_html(self, result_df, multiple_gtips):
        """GTİP-Tanım Detay Analizi için HTML raporu oluştur - Daha anlaşılır format"""
        if result_df.empty:
            return "<p>Aynı ticari tanımda farklı GTİP kodu kullanımı tespit edilmedi.</p>"
        
        # İstatistikleri hesapla
        unique_ticari_tanim = len(multiple_gtips)
        unique_gtip = result_df['Gtip'].nunique()
        unique_firma = result_df['Adi_unvani'].nunique()
        total_records = len(result_df)
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>GTİP-Tanım Detay Analizi</title>
            <style>
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    margin: 0;
                    padding: 20px;
                    background-color: #f8f9fa;
                    color: #333;
                    line-height: 1.6;
                }}
                .container {{
                    max-width: 1600px;
                    margin: 0 auto;
                    background-color: white;
                    border-radius: 12px;
                    box-shadow: 0 4px 20px rgba(0,0,0,0.1);
                    overflow: hidden;
                }}
                .header {{
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 30px;
                    text-align: center;
                }}
                .header h1 {{
                    margin: 0;
                    font-size: 28px;
                    font-weight: 300;
                    margin-bottom: 10px;
                }}
                .header .subtitle {{
                    font-size: 16px;
                    opacity: 0.9;
                }}
                .stats-row {{
                    display: flex;
                    justify-content: space-around;
                    padding: 30px;
                    background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
                    flex-wrap: wrap;
                }}
                .stat-card {{
                    background: white;
                    border-radius: 10px;
                    padding: 20px;
                    text-align: center;
                    min-width: 160px;
                    margin: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                    border-left: 4px solid #667eea;
                }}
                .stat-number {{
                    font-size: 28px;
                    font-weight: bold;
                    color: #667eea;
                    display: block;
                    margin-bottom: 5px;
                }}
                .stat-label {{
                    font-size: 14px;
                    color: #6c757d;
                    font-weight: 500;
                }}
                .content {{
                    padding: 30px;
                }}
                .section-title {{
                    font-size: 20px;
                    font-weight: bold;
                    color: #0d47a1;
                    margin: 20px 0;
                    padding: 15px;
                    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
                    border-radius: 8px;
                    border-left: 4px solid #2196f3;
                }}
                .tanim-block {{
                    margin-bottom: 40px;
                    border: 1px solid #dee2e6;
                    border-radius: 12px;
                    overflow: hidden;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.05);
                }}
                .tanim-header {{
                    background: linear-gradient(135deg, #fff3e0 0%, #ffe0b2 100%);
                    padding: 20px;
                    border-bottom: 1px solid #ffb74d;
                }}
                .tanim-title {{
                    font-size: 16px;
                    font-weight: bold;
                    color: #e65100;
                    margin: 0 0 10px 0;
                    word-wrap: break-word;
                }}
                .summary-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 0;
                    background: white;
                }}
                .summary-table th {{
                    background: linear-gradient(135deg, #1565c0 0%, #1976d2 100%);
                    color: white;
                    padding: 15px 12px;
                    text-align: left;
                    font-weight: 600;
                    font-size: 14px;
                }}
                .summary-table td {{
                    padding: 12px;
                    border-bottom: 1px solid #e0e0e0;
                    vertical-align: top;
                    font-size: 13px;
                }}
                .summary-table tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                .summary-table tr:hover {{
                    background-color: #e3f2fd;
                    transition: background-color 0.3s ease;
                }}
                .gtip-code {{
                    font-family: 'Courier New', monospace;
                    font-weight: bold;
                    color: #d32f2f;
                    background-color: #ffebee;
                    padding: 6px 12px;
                    border-radius: 6px;
                    border: 1px solid #ffcdd2;
                    display: inline-block;
                    margin: 2px;
                    font-size: 12px;
                }}
                .beyanname-count {{
                    background: linear-gradient(135deg, #e8f5e8 0%, #c8e6c9 100%);
                    color: #2e7d32;
                    padding: 6px 16px;
                    border-radius: 20px;
                    font-size: 12px;
                    font-weight: bold;
                    border: 1px solid #81c784;
                    display: inline-block;
                    text-align: center;
                }}
                .detail-btn {{
                    background: #007bff;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    cursor: pointer;
                    font-size: 12px;
                    margin: 15px 0;
                    transition: all 0.3s ease;
                    font-weight: 500;
                }}
                .detail-btn:hover {{
                    background: #0056b3;
                    transform: translateY(-1px);
                }}
                .detail-btn.expanded {{
                    background: #28a745;
                }}
                .detail-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 15px;
                    display: none;
                    background: #fafafa;
                }}
                .detail-table th {{
                    background: linear-gradient(135deg, #37474f 0%, #455a64 100%);
                    color: white;
                    padding: 12px 10px;
                    text-align: left;
                    font-weight: 600;
                    font-size: 12px;
                }}
                .detail-table td {{
                    padding: 10px;
                    border-bottom: 1px solid #e0e0e0;
                    font-size: 11px;
                }}
                .detail-table tr:nth-child(even) {{
                    background-color: #f5f5f5;
                }}
                .detail-table tr:hover {{
                    background-color: #e1f5fe;
                }}
                .firma-cell {{
                    max-width: 250px;
                    word-wrap: break-word;
                    font-weight: 500;
                    color: #424242;
                }}
                .beyanname-cell {{
                    background: #17a2b8;
                    color: white;
                    padding: 4px 10px;
                    border-radius: 12px;
                    font-size: 10px;
                    font-weight: 500;
                    text-align: center;
                    display: inline-block;
                }}
                .highlight-red {{ background-color: #ffebee; }}
                .highlight-blue {{ background-color: #e3f2fd; }}
                .highlight-green {{ background-color: #e8f5e8; }}
                .highlight-yellow {{ background-color: #fff3e0; }}
            </style>
            <script>
                function toggleDetails(tableId, btnId) {{
                    var table = document.getElementById(tableId);
                    var button = document.getElementById(btnId);
                    
                    if (table.style.display === 'none' || table.style.display === '') {{
                        table.style.display = 'table';
                        button.innerHTML = '🔼 Detayları Gizle';
                        button.classList.add('expanded');
                    }} else {{
                        table.style.display = 'none';
                        button.innerHTML = '🔽 Detayları Göster';
                        button.classList.remove('expanded');
                    }}
                }}
            </script>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>GTİP-Tanım Detay Analizi</h1>
                    <p class="subtitle">Aynı ticari tanımda farklı GTİP kodları kullanılan eşyaların detaylı analizi</p>
                </div>
                
                <div class="stats-row">
                    <div class="stat-card">
                        <span class="stat-number">{unique_ticari_tanim}</span>
                        <div class="stat-label">Tutarsız Ticari Tanım</div>
                    </div>
                    <div class="stat-card">
                        <span class="stat-number">{unique_gtip}</span>
                        <div class="stat-label">Farklı GTİP Kodu</div>
                    </div>
                    <div class="stat-card">
                        <span class="stat-number">{unique_firma}</span>
                        <div class="stat-label">Farklı Firma</div>
                    </div>
                    <div class="stat-card">
                        <span class="stat-number">{total_records}</span>
                        <div class="stat-label">Toplam Kayıt</div>
                    </div>
                </div>
                
                <div class="content">
                    <div class="section-title">🔍 Detaylı Analiz Sonuçları</div>
        """
        
        # Her ticari tanım için gruplu analiz
        colors = ['highlight-red', 'highlight-blue', 'highlight-green', 'highlight-yellow']
        
        for idx, (_, tanim_row) in enumerate(multiple_gtips.iterrows()):
            ticari_tanim = tanim_row['Ticari_tanimi']
            
            # Bu ticari tanıma ait kayıtları filtrele
            tanim_records = result_df[result_df['Ticari_tanimi'] == ticari_tanim]
            
            # GTİP kodlarını grupla ve sayıları hesapla
            gtip_analysis = tanim_records.groupby('Gtip').agg({
                'Beyanname_no': 'nunique',
                'Adi_unvani': lambda x: list(x.unique())
            }).reset_index()
            gtip_analysis.columns = ['Gtip', 'Beyanname_Sayisi', 'Firmalar']
            gtip_analysis['Firma_Sayisi'] = gtip_analysis['Firmalar'].apply(len)
            gtip_analysis = gtip_analysis.sort_values('Beyanname_Sayisi', ascending=False)
            
            table_id = f"detail-table-{idx}"
            btn_id = f"detail-btn-{idx}"
            color_class = colors[idx % len(colors)]
            
            html += f"""
                    <div class="tanim-block {color_class}">
                        <div class="tanim-header">
                            <div class="tanim-title">📦 {ticari_tanim}</div>
                            <div style="margin-top: 10px; font-size: 13px; color: #666;">
                                <strong>{len(gtip_analysis)} farklı GTİP kodu</strong> - 
                                <strong>{tanim_records['Beyanname_no'].nunique()} beyanname</strong> - 
                                <strong>{tanim_records['Adi_unvani'].nunique()} firma</strong>
                            </div>
                        </div>
                        
                        <table class="summary-table">
                            <thead>
                                <tr>
                                    <th style="width: 20%;">GTİP Kodu</th>
                                    <th style="width: 15%;">Beyanname Sayısı</th>
                                    <th style="width: 15%;">Firma Sayısı</th>
                                    <th style="width: 50%;">Kullanılan Firmalar</th>
                                </tr>
                            </thead>
                            <tbody>
            """
            
            # GTİP kodları için özet satırlar
            for _, gtip_row in gtip_analysis.iterrows():
                gtip_code = gtip_row['Gtip']
                beyanname_count = gtip_row['Beyanname_Sayisi']
                firma_count = gtip_row['Firma_Sayisi']
                firmalar = gtip_row['Firmalar']
                
                # Firma listesini kısalt ve daha okunabilir yap
                if len(firmalar) <= 2:
                    firma_display = ', '.join(firmalar)
                else:
                    firma_display = f"{firmalar[0]}, {firmalar[1]} + {len(firmalar)-2} firma daha"
                
                html += f"""
                                <tr>
                                    <td><span class="gtip-code">{gtip_code}</span></td>
                                    <td><span class="beyanname-count">{beyanname_count} beyanname</span></td>
                                    <td><span class="beyanname-count">{firma_count} firma</span></td>
                                    <td><div class="firma-cell">{firma_display}</div></td>
                                </tr>
                """
            
            html += f"""
                            </tbody>
                        </table>
                        
                        <button class="detail-btn" id="{btn_id}" onclick="toggleDetails('{table_id}', '{btn_id}')">
                            🔽 Detayları Göster
                        </button>
                        
                        <table class="detail-table" id="{table_id}">
                            <thead>
                                <tr>
                                    <th style="width: 15%;">GTİP Kodu</th>
                                    <th style="width: 30%;">Firma Unvanı</th>
                                    <th style="width: 15%;">Beyanname No</th>
                                    <th style="width: 10%;">Kalem No</th>
                                    <th style="width: 12%;">Menşei Ülke</th>
                                    <th style="width: 18%;">Fatura Bilgisi</th>
                                </tr>
                            </thead>
                            <tbody>
            """
            
            # Detaylı kayıtları GTİP koduna göre sırala
            sorted_records = tanim_records.sort_values(['Gtip', 'Adi_unvani', 'Beyanname_no'])
            
            for _, record in sorted_records.iterrows():
                gtip = record.get('Gtip', '')
                firma = record.get('Adi_unvani', '')
                beyanname = record.get('Beyanname_no', '')
                kalem_no = record.get('BeyannameKalemNo', '')
                mensei = record.get('Mensei_ulke', '')
                fatura_miktar = record.get('Fatura_miktari', '')
                fatura_doviz = record.get('Fatura_miktarinin_dovizi', '')
                
                fatura_info = f"{fatura_miktar} {fatura_doviz}" if fatura_miktar and fatura_doviz else "-"
                
                html += f"""
                                <tr>
                                    <td><span class="gtip-code">{gtip}</span></td>
                                    <td><div class="firma-cell">{firma}</div></td>
                                    <td><span class="beyanname-cell">{beyanname}</span></td>
                                    <td>{kalem_no or '-'}</td>
                                    <td>{mensei or '-'}</td>
                                    <td>{fatura_info}</td>
                                </tr>
                """
            
            html += """
                            </tbody>
                        </table>
                    </div>
            """
        
        html += """
                </div>
            </div>
        </body>
        </html>
        """
        
        return html

    def check_rarely_used_origin_country_by_sender_gtip(self):
        """Aynı gönderici ve GTİP kodunda nadiren kullanılan menşe ülke kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Gönderici-GTİP bazında nadir menşe ülke analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Analiz fonksiyonunu çağır
            sender_gtip_check = check_rarely_used_origin_country_by_sender_gtip(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if sender_gtip_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", sender_gtip_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(sender_gtip_check["message"])
                return
            
            result = {
                "Gönderici-GTİP Bazında Nadir Menşe Ülke Analizi": sender_gtip_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in sender_gtip_check:
                self.check_results_widget.set_html_report(sender_gtip_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in sender_gtip_check and "summary" in sender_gtip_check:
                self.check_results_widget.show_details(sender_gtip_check["data"])
                self.check_results_widget.show_summary(sender_gtip_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(sender_gtip_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Gönderici-GTİP bazında nadir menşe ülke analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def run_all_analyses_and_export(self):
        """Tüm analizleri çalıştır ve sonuçları tek Excel dosyasına aktar"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Önce veri yükleyin!")
            return
        
        # Progress dialog oluştur
        progress_dialog = QDialog(self)
        progress_dialog.setWindowTitle("Tüm Analizler Çalıştırılıyor...")
        progress_dialog.setModal(True)
        progress_dialog.setFixedSize(400, 150)
        
        layout = QVBoxLayout(progress_dialog)
        
        progress_label = QLabel("Analizler başlatılıyor...")
        layout.addWidget(progress_label)
        
        progress_bar = QProgressBar()
        progress_bar.setRange(0, 100)
        layout.addWidget(progress_bar)
        
        cancel_btn = QPushButton("İptal")
        layout.addWidget(cancel_btn)
        
        # İptal flag'i
        self.analysis_cancelled = False
        cancel_btn.clicked.connect(lambda: setattr(self, 'analysis_cancelled', True))
        
        progress_dialog.show()
        QApplication.processEvents()
        
        try:
            # Tüm analiz fonksiyonları listesi
            analysis_functions = [
                ("Eksik Değerler", self.check_missing_values),
                ("Tekrarlanan Satırlar", self.check_duplicate_rows),
                ("Ağırlık Tutarlılığı", self.check_weight_consistency),
                ("İşlem Niteliği Tutarlılığı", self.check_islem_niteligi_consistency),
                ("GTİP-Ürün Kodu Tutarlılığı", self.check_gtip_urun_kodu_consistency),
                ("Alıcı-Satıcı İlişkisi", self.check_alici_satici_relationship),
                ("GTİP Tanım Detayı", self.check_gtip_tanim_detail),
                ("Nadir Döviz Kullanımı", self.check_rarely_used_currency),
                ("Nadir Menşe Ülke", self.check_rarely_used_origin_country),
                ("Gönderici-GTİP Nadir Menşe", self.check_rarely_used_origin_country_by_sender_gtip),
                ("Nadir Ödeme Şekli", self.check_rarely_used_payment_method),
                ("Birim Fiyat Artışı", self.check_unit_price_increase),
                ("KDV Kontrol", self.check_kdv_consistency),
                ("Yurt İçi Gider Değişimi", self.check_domestic_expense_variation),
                ("Yurt Dışı Gider Değişimi", self.check_foreign_expense_variation),
                ("Supalan Depolama Kontrol", self.check_supalan_storage),
                ("IGV Kontrol", self.check_igv_consistency),
                ("Tedarikçi Beyan Kontrol", self.check_tedarikci_beyan_kontrol)
            ]
            
            total_analyses = len(analysis_functions)
            all_results = {}
            
            # Her analizi çalıştır
            for i, (analysis_name, analysis_func) in enumerate(analysis_functions):
                if self.analysis_cancelled:
                    progress_dialog.close()
                    return
                
                progress_label.setText(f"Çalıştırılıyor: {analysis_name}")
                progress_bar.setValue(int((i / total_analyses) * 90))  # %90'a kadar analiz
                QApplication.processEvents()
                
                try:
                    # Analizi çalıştır
                    analysis_func()
                    
                    # Kısa bir bekleme süresi ekle (analiz tamamlanması için)
                    QApplication.processEvents()
                    time.sleep(0.1)
                    
                    # Sonucu al - daha geniş arama
                    if hasattr(self.check_results_widget, 'check_results'):
                        # Önce tam eşleşme ara
                        for check_name, result in self.check_results_widget.check_results.items():
                            if check_name == analysis_name:
                                all_results[analysis_name] = result
                                break
                        else:
                            # Tam eşleşme bulunamazsa kısmi eşleşme ara
                            for check_name, result in self.check_results_widget.check_results.items():
                                # Anahtar kelimeler ile eşleştir
                                analysis_keywords = analysis_name.lower().split()
                                check_keywords = check_name.lower().split()
                                
                                # En az bir anahtar kelime eşleşirse
                                if any(keyword in check_name.lower() for keyword in analysis_keywords) or \
                                   any(keyword in analysis_name.lower() for keyword in check_keywords):
                                    all_results[analysis_name] = result
                                    break
                    
                except Exception as e:
                    print(f"Analiz hatası ({analysis_name}): {str(e)}")
                    continue
            
            if self.analysis_cancelled:
                progress_dialog.close()
                return
            
            # Excel'e aktar
            progress_label.setText("Excel dosyası oluşturuluyor...")
            progress_bar.setValue(95)
            QApplication.processEvents()
            
            # Dosya kaydetme dialog'u
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Tüm Analiz Sonuçlarını Kaydet", 
                f"tum_analiz_sonuclari_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "Excel Files (*.xlsx)"
            )
            
            if file_path:
                if not file_path.lower().endswith('.xlsx'):
                    file_path += '.xlsx'
                
                # Excel dosyası oluştur
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    # Tüm analiz sonuçlarını birleştirmek için liste
                    all_combined_data = []
                    sheet_index = 1
                    
                    # Önce her analiz için ayrı sheet'ler oluştur
                    for analysis_name, result in all_results.items():
                        if result and "data" in result and result["data"] is not None:
                            df_to_export = result["data"]
                            if isinstance(df_to_export, pd.DataFrame) and not df_to_export.empty:
                                # Ayrı sheet oluştur
                                sheet_name = analysis_name[:25] + f"_{sheet_index}"
                                sheet_name = sheet_name.replace("/", "_").replace("\\", "_")
                                df_to_export.to_excel(writer, sheet_name=sheet_name, index=False)
                                
                                # Birleştirme için de hazırla
                                df_copy = df_to_export.copy()
                                df_copy.insert(0, 'Analiz_Türü', analysis_name)
                                all_combined_data.append(df_copy)
                                
                                sheet_index += 1
                    
                    # Şimdi tüm verileri birleştiren sheet'i oluştur
                    if all_combined_data:
                        # Tüm DataFrame'leri alt alta birleştir
                        combined_df = pd.concat(all_combined_data, ignore_index=True, sort=False)
                        
                        # Birleştirilmiş sheet'e yaz
                        combined_df.to_excel(writer, sheet_name='Tüm_Analiz_Sonuçları', index=False)
                        
                        # Birleştirilmiş sheet'i formatla
                        workbook = writer.book
                        worksheet = writer.sheets['Tüm_Analiz_Sonuçları']
                        
                        # Başlık satırını formatla
                        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                        header_font = Font(bold=True, color="FFFFFF", size=12)
                        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                        
                        # Kenarlık stili
                        thin_border = Border(
                            left=Side(style='thin'),
                            right=Side(style='thin'),
                            top=Side(style='thin'),
                            bottom=Side(style='thin')
                        )
                        
                        # Başlık satırını formatla
                        for col_num in range(1, len(combined_df.columns) + 1):
                            cell = worksheet.cell(row=1, column=col_num)
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                            cell.border = thin_border
                        
                        # Analiz türü sütununa göre satırları renklendir
                        current_analysis = None
                        color_index = 0
                        colors = [
                            "F8F9FA",  # Açık gri
                            "E3F2FD",  # Açık mavi
                            "F3E5F5",  # Açık mor
                            "E8F5E8",  # Açık yeşil
                            "FFF3E0",  # Açık turuncu
                            "FFEBEE",  # Açık kırmızı
                        ]
                        
                        for row_num in range(2, len(combined_df) + 2):
                            analysis_type = worksheet.cell(row=row_num, column=1).value
                            
                            # Yeni analiz türü başladığında renk değiştir
                            if analysis_type != current_analysis:
                                current_analysis = analysis_type
                                color_index = (color_index + 1) % len(colors)
                            
                            # Satırı renklendir
                            row_fill = PatternFill(start_color=colors[color_index], end_color=colors[color_index], fill_type="solid")
                            
                            for col_num in range(1, len(combined_df.columns) + 1):
                                cell = worksheet.cell(row=row_num, column=col_num)
                                cell.fill = row_fill
                                cell.border = thin_border
                                cell.alignment = Alignment(horizontal="left", vertical="center")
                        
                        # Sütun genişliklerini otomatik ayarla
                        for column in worksheet.columns:
                            max_length = 0
                            column_letter = column[0].column_letter
                            
                            for cell in column:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(str(cell.value))
                                except:
                                    pass
                            
                            # Maksimum genişlik sınırı
                            adjusted_width = min(max_length + 2, 50)
                            worksheet.column_dimensions[column_letter].width = adjusted_width
                        
                        # Başlık satırının yüksekliğini ayarla
                        worksheet.row_dimensions[1].height = 25
                
                progress_bar.setValue(100)
                progress_label.setText("Tamamlandı!")
                QApplication.processEvents()
                
                # Başarı mesajı
                QMessageBox.information(
                    self, "Başarılı", 
                    f"Tüm analiz sonuçları başarıyla kaydedildi:\n{file_path}\n\n"
                    f"📊 Toplam {len(all_results)} analiz sonucu aktarıldı.\n"
                    f"📁 Her analiz için ayrı sheet + 1 birleştirilmiş sheet oluşturuldu.\n"
                    f"🔗 'Tüm_Analiz_Sonuçları' sheet'inde tüm veriler birleştirilmiş halde."
                )
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Analiz sırasında hata oluştu:\n{str(e)}")
        
        finally:
            progress_dialog.close()

    def create_word_report(self):
        """Tüm görsel raporları Word belgesine aktar"""
        if not hasattr(self.check_results_widget, 'check_results') or not self.check_results_widget.check_results:
            QMessageBox.warning(self, "Uyarı", "Önce analizleri çalıştırın!")
            return
        
        # Progress dialog oluştur
        progress_dialog = QDialog(self)
        progress_dialog.setWindowTitle("Word Raporu Oluşturuluyor...")
        progress_dialog.setModal(True)
        progress_dialog.setFixedSize(400, 150)
        
        layout = QVBoxLayout(progress_dialog)
        
        progress_label = QLabel("Word belgesi hazırlanıyor...")
        layout.addWidget(progress_label)
        
        progress_bar = QProgressBar()
        progress_bar.setRange(0, 100)
        layout.addWidget(progress_bar)
        
        cancel_btn = QPushButton("İptal")
        layout.addWidget(cancel_btn)
        
        # İptal flag'i
        self.word_cancelled = False
        cancel_btn.clicked.connect(lambda: setattr(self, 'word_cancelled', True))
        
        progress_dialog.show()
        QApplication.processEvents()
        
        try:
            # Dosya kaydetme dialog'u
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Word Raporu Kaydet", 
                f"beyanname_analiz_raporu_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx)"
            )
            
            if not file_path:
                progress_dialog.close()
                return
            
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'
            
            progress_label.setText("Word kütüphanesi yükleniyor...")
            progress_bar.setValue(10)
            QApplication.processEvents()
            
            # python-docx kütüphanesini import et
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT
                from docx.oxml.shared import OxmlElement, qn
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                import re
                from bs4 import BeautifulSoup
                import tempfile
                import os
            except ImportError as e:
                QMessageBox.critical(
                    self, "Hata", 
                    f"Word raporu oluşturmak için gerekli kütüphaneler eksik!\n\n"
                    f"Hata: {str(e)}\n\n"
                    f"Lütfen şu komutları çalıştırın:\n"
                    f"pip install python-docx beautifulsoup4"
                )
                progress_dialog.close()
                return
            
            if self.word_cancelled:
                progress_dialog.close()
                return
            
            progress_label.setText("Word belgesi oluşturuluyor...")
            progress_bar.setValue(20)
            QApplication.processEvents()
            
            # Word belgesi oluştur
            doc = Document()
            
            # Sayfa kenar boşluklarını ayarla
            try:
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.8)
                    section.right_margin = Inches(0.8)
            except:
                pass  # Kenar boşluğu ayarlanamadıysa devam et
            
            # KAPAK SAYFASI
            # Ana başlık
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run('🏛️ BEYANNAME ANALİZ RAPORU')
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(52, 73, 94)
            
            # Alt başlık
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run('Gümrük Beyannamesi Tutarlılık ve Risk Analizi')
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.color.rgb = RGBColor(127, 140, 141)
            
            # Tarih
            doc.add_paragraph()
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'📅 Rapor Tarihi: {pd.Timestamp.now().strftime("%d.%m.%Y %H:%M")}')
            date_run.font.size = Pt(14)
            date_run.font.color.rgb = RGBColor(52, 152, 219)
            
            # DASHBOARD OLUŞTUR
            doc.add_paragraph()
            doc.add_paragraph()
            
            dashboard_title = doc.add_heading('📊 YÖNETİCİ ÖZETİ', level=1)
            dashboard_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Analiz istatistiklerini topla
            total_analyses = len(self.check_results_widget.check_results)
            error_count = 0
            warning_count = 0
            success_count = 0
            total_records = 0
            
            for result in self.check_results_widget.check_results.values():
                if result:
                    status = result.get('status', 'unknown')
                    if status == 'error':
                        error_count += 1
                    elif status == 'warning':
                        warning_count += 1
                    elif status == 'success':
                        success_count += 1
                    
                    if 'data' in result and result['data'] is not None:
                        try:
                            total_records += len(result['data'])
                        except:
                            pass
            
            # Dashboard tablosu oluştur
            dashboard_table = doc.add_table(rows=5, cols=4)
            dashboard_table.style = 'Table Grid'
            dashboard_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Başlık satırı
            header_cells = dashboard_table.rows[0].cells
            headers = ['📈 GENEL İSTATİSTİKLER', '🔍 TOPLAM ANALİZ', '📊 İNCELENEN KAYIT', '⚡ DURUM']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Başlık formatı
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Mavi arka plan
                self._set_cell_background_color(header_cells[i], RGBColor(52, 152, 219))
            
            # Veri satırları
            data_rows = [
                ['🔴 Kritik Hatalar', str(error_count), f'{(error_count/total_analyses*100):.1f}%', 'ACİL MÜDAHALE'],
                ['🟡 Uyarılar', str(warning_count), f'{(warning_count/total_analyses*100):.1f}%', 'İNCELENMELİ'],
                ['🟢 Başarılı', str(success_count), f'{(success_count/total_analyses*100):.1f}%', 'SORUNSUZ'],
                ['📋 TOPLAM', str(total_analyses), f'{total_records:,}', 'TAMAMLANDI']
            ]
            
            colors = [RGBColor(231, 76, 60), RGBColor(243, 156, 18), RGBColor(39, 174, 96), RGBColor(52, 73, 94)]
            
            for i, row_data in enumerate(data_rows, 1):
                cells = dashboard_table.rows[i].cells
                for j, cell_data in enumerate(row_data):
                    cells[j].text = cell_data
                    cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Renk kodlaması
                    if j == 0:  # İlk sütun
                        for paragraph in cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = colors[i-1]
                                if i == 4:  # Toplam satırı
                                    run.font.bold = True
                    elif i == 4:  # Toplam satırı
                        for paragraph in cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            # Sayfa sonu
            doc.add_page_break()
            
            # DETAY ANALİZLER
            detail_title = doc.add_heading('📋 DETAY ANALİZ SONUÇLARI', level=1)
            
            # Tüm analiz sonuçlarını topla (HTML raporu olan ve olmayan)
            all_results = []
            for check_name, result in self.check_results_widget.check_results.items():
                if result:
                    all_results.append((check_name, result))
            
            total_reports = len(all_results)
            
            if total_reports == 0:
                doc.add_paragraph("Analiz sonucu bulunamadı.")
            else:
                # Geçici dizin oluştur (resimler için)
                temp_dir = tempfile.mkdtemp()
                
                # Her analiz için rapor ekle
                for i, (check_name, result) in enumerate(all_results):
                    if self.word_cancelled:
                        progress_dialog.close()
                        return
                    
                    progress_label.setText(f"İşleniyor: {check_name}")
                    progress_bar.setValue(30 + int((i / total_reports) * 60))
                    QApplication.processEvents()
                    
                    # Analiz başlığı
                    heading = doc.add_heading(f'{i+1}. {check_name.upper()}', level=2)
                    heading_run = heading.runs[0]
                    heading_run.font.color.rgb = RGBColor(52, 73, 94)
                    
                    # Durum bilgisi tablosu
                    status = result.get('status', 'unknown')
                    message = result.get('message', 'Bilgi yok')
                    
                    status_table = doc.add_table(rows=1, cols=2)
                    status_table.style = 'Table Grid'
                    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    status_cells = status_table.rows[0].cells
                    status_cells[0].text = f'📊 DURUM: {status.upper()}'
                    status_cells[1].text = f'💬 SONUÇ: {message}'
                    
                    # Durum renk kodlaması
                    status_color = RGBColor(39, 174, 96) if status == 'success' else RGBColor(231, 76, 60) if status == 'error' else RGBColor(243, 156, 18)
                    
                    for cell in status_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = status_color
                    
                    doc.add_paragraph()  # Boşluk
                    
                    # HTML raporu varsa görsel olarak yakala, yoksa veri tablosu ekle
                    if "html_report" in result and result["html_report"]:
                        # HTML raporunu görsel olarak yakala ve Word'e ekle
                        html_content = result["html_report"]
                        try:
                            # HTML'i görsel olarak yakala
                            image_path = self._capture_html_as_image(html_content, temp_dir, f"report_{i}")
                            if image_path and os.path.exists(image_path):
                                # Resmi Word'e ekle
                                para = doc.add_paragraph()
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.runs[0] if para.runs else para.add_run()
                                
                                # Resim boyutunu ayarla (sayfa genişliğine sığacak şekilde)
                                try:
                                    run.add_picture(image_path, width=Inches(6.5))
                                except:
                                    # Resim eklenemezse HTML'i metin olarak ekle
                                    self._convert_html_to_word(doc, html_content)
                            else:
                                # Görsel yakalanamadıysa HTML'i metin olarak ekle
                                self._convert_html_to_word(doc, html_content)
                                
                        except Exception as e:
                            # Hata durumunda HTML'i metin olarak ekle
                            print(f"HTML görsel yakalama hatası: {str(e)}")
                            try:
                                self._convert_html_to_word(doc, html_content)
                            except Exception as e2:
                                error_para = doc.add_paragraph(f"⚠️ Rapor içeriği işlenirken hata oluştu: {str(e2)}")
                                error_run = error_para.runs[0]
                                error_run.font.color.rgb = RGBColor(231, 76, 60)
                    else:
                        # HTML raporu yoksa veri tablosu ekle
                        if 'data' in result and result['data'] is not None:
                            try:
                                data = result['data']
                                if hasattr(data, 'head'):  # DataFrame ise
                                    # DataFrame'i Word tablosuna çevir
                                    if len(data) > 0:
                                        # En fazla 20 satır göster
                                        display_data = data.head(20)
                                        
                                        # Tablo oluştur
                                        table = doc.add_table(rows=len(display_data) + 1, cols=len(display_data.columns))
                                        table.style = 'Table Grid'
                                        
                                        # Başlık satırı
                                        header_cells = table.rows[0].cells
                                        for j, col_name in enumerate(display_data.columns):
                                            header_cells[j].text = str(col_name)
                                            for paragraph in header_cells[j].paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.bold = True
                                        
                                        # Veri satırları
                                        for row_idx, (_, row) in enumerate(display_data.iterrows(), 1):
                                            cells = table.rows[row_idx].cells
                                            for col_idx, value in enumerate(row):
                                                cells[col_idx].text = str(value) if pd.notna(value) else ""
                                        
                                        if len(data) > 20:
                                            doc.add_paragraph(f"Not: Toplam {len(data)} kayıt bulundu, ilk 20 tanesi gösterildi.")
                                    else:
                                        doc.add_paragraph("Veri bulunamadı.")
                                else:
                                    # DataFrame değilse metin olarak ekle
                                    doc.add_paragraph(str(data))
                            except Exception as e:
                                error_para = doc.add_paragraph(f"⚠️ Veri tablosu oluşturulurken hata: {str(e)}")
                                error_run = error_para.runs[0]
                                error_run.font.color.rgb = RGBColor(231, 76, 60)
                        else:
                            doc.add_paragraph("Bu analiz için detay veri bulunmuyor.")
                    
                    # Sayfa sonu (son analiz değilse)
                    if i < total_reports - 1:
                        doc.add_page_break()
                
                # Geçici dosyaları temizle
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass
            
            if self.word_cancelled:
                progress_dialog.close()
                return
            
            progress_label.setText("Word belgesi kaydediliyor...")
            progress_bar.setValue(95)
            QApplication.processEvents()
            
            # Belgeyi kaydet
            doc.save(file_path)
            
            progress_bar.setValue(100)
            progress_label.setText("Tamamlandı!")
            QApplication.processEvents()
            
            # Başarı mesajı
            QMessageBox.information(
                self, "Başarılı", 
                f"📄 Word raporu başarıyla oluşturuldu:\n{file_path}\n\n"
                f"📊 Dashboard + {total_reports} detay analiz raporu\n"
                f"🎯 Yönetici sunumu için hazır!"
            )
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Word raporu oluşturulurken hata:\n{str(e)}")
        
        finally:
            progress_dialog.close()
    
    def _capture_html_as_image(self, html_content, temp_dir, filename):
        """HTML içeriğini görsel olarak yakala ve resim dosyası olarak kaydet"""
        try:
            # Yöntem 1: wkhtmltoimage (en güvenilir)
            image_path = self._capture_with_wkhtmltoimage(html_content, temp_dir, filename)
            if image_path:
                return image_path
            
            # Yöntem 2: Playwright (modern ve güçlü)
            image_path = self._capture_with_playwright(html_content, temp_dir, filename)
            if image_path:
                return image_path
            
            # Yöntem 3: WeasyPrint (pure Python)
            image_path = self._capture_with_weasyprint(html_content, temp_dir, filename)
            if image_path:
                return image_path
            
            # Yöntem 4: QWebEngineView (fallback)
            return self._capture_with_webengine(html_content, temp_dir, filename)
            
        except Exception as e:
            print(f"HTML görsel yakalama hatası: {str(e)}")
            return None
    
    def _capture_with_wkhtmltoimage(self, html_content, temp_dir, filename):
        """wkhtmltoimage ile HTML görsel yakalama (en güvenilir yöntem)"""
        try:
            import subprocess
            import os
            import shutil
            
            # wkhtmltoimage'ın yüklü olup olmadığını kontrol et
            if not shutil.which('wkhtmltoimage'):
                print("wkhtmltoimage bulunamadı. Yükleme talimatları:")
                print("Windows: https://wkhtmltopdf.org/downloads.html adresinden indirin")
                print("Linux: sudo apt-get install wkhtmltopdf")
                print("macOS: brew install wkhtmltopdf")
                return None
            
            # Geçici HTML dosyası oluştur
            html_file_path = os.path.join(temp_dir, f"{filename}.html")
            optimized_html = self._optimize_html_for_capture(html_content)
            
            with open(html_file_path, 'w', encoding='utf-8') as f:
                f.write(optimized_html)
            
            # Çıktı dosyası
            image_path = os.path.join(temp_dir, f"{filename}.png")
            
            # wkhtmltoimage komutunu çalıştır
            cmd = [
                'wkhtmltoimage',
                '--width', '1200',
                '--height', '1600',
                '--quality', '90',
                '--format', 'png',
                '--enable-local-file-access',
                '--javascript-delay', '1000',
                '--no-stop-slow-scripts',
                '--debug-javascript',
                html_file_path,
                image_path
            ]
            
            # Komutu çalıştır
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and os.path.exists(image_path):
                print(f"wkhtmltoimage ile başarılı: {image_path}")
                return image_path
            else:
                print(f"wkhtmltoimage hatası: {result.stderr}")
                return None
                
        except subprocess.TimeoutExpired:
            print("wkhtmltoimage timeout")
            return None
        except Exception as e:
            print(f"wkhtmltoimage hatası: {str(e)}")
            return None
    
    def _capture_with_playwright(self, html_content, temp_dir, filename):
        """Playwright ile HTML görsel yakalama"""
        try:
            # Playwright'ı import et
            from playwright.sync_api import sync_playwright
            import os
            
            # Geçici HTML dosyası oluştur
            html_file_path = os.path.join(temp_dir, f"{filename}.html")
            optimized_html = self._optimize_html_for_capture(html_content)
            
            with open(html_file_path, 'w', encoding='utf-8') as f:
                f.write(optimized_html)
            
            # Çıktı dosyası
            image_path = os.path.join(temp_dir, f"{filename}.png")
            
            # Playwright ile ekran görüntüsü al
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()
                
                # Sayfa boyutunu ayarla
                page.set_viewport_size({"width": 1200, "height": 1600})
                
                # HTML dosyasını yükle
                page.goto(f"file://{html_file_path}")
                
                # Sayfa yüklenene kadar bekle
                page.wait_for_load_state('networkidle')
                
                # Ekran görüntüsü al
                page.screenshot(path=image_path, full_page=True)
                
                browser.close()
            
            if os.path.exists(image_path):
                print(f"Playwright ile başarılı: {image_path}")
                return image_path
            else:
                return None
                
        except ImportError:
            print("Playwright bulunamadı. Yüklemek için: pip install playwright && playwright install")
            return None
        except Exception as e:
            print(f"Playwright hatası: {str(e)}")
            return None
    
    def _capture_with_weasyprint(self, html_content, temp_dir, filename):
        """WeasyPrint ile HTML'i PNG'ye çevir"""
        try:
            from weasyprint import HTML, CSS
            from PIL import Image
            import os
            import io
            
            # Geçici HTML dosyası oluştur
            html_file_path = os.path.join(temp_dir, f"{filename}.html")
            optimized_html = self._optimize_html_for_capture(html_content)
            
            with open(html_file_path, 'w', encoding='utf-8') as f:
                f.write(optimized_html)
            
            # Önce PDF oluştur
            pdf_path = os.path.join(temp_dir, f"{filename}.pdf")
            HTML(filename=html_file_path).write_pdf(pdf_path)
            
            # PDF'i PNG'ye çevir (pdf2image gerekli)
            try:
                from pdf2image import convert_from_path
                
                # PDF'i resme çevir
                images = convert_from_path(pdf_path, dpi=150)
                if images:
                    image_path = os.path.join(temp_dir, f"{filename}.png")
                    images[0].save(image_path, 'PNG')
                    
                    # PDF dosyasını sil
                    os.remove(pdf_path)
                    
                    print(f"WeasyPrint ile başarılı: {image_path}")
                    return image_path
                    
            except ImportError:
                print("pdf2image bulunamadı. WeasyPrint sadece PDF oluşturabilir.")
                return None
                
        except ImportError:
            print("WeasyPrint bulunamadı. Yüklemek için: pip install weasyprint")
            return None
        except Exception as e:
            print(f"WeasyPrint hatası: {str(e)}")
            return None
    
    def _capture_with_webengine(self, html_content, temp_dir, filename):
        """QWebEngineView ile HTML görsel yakalama (fallback yöntem)"""
        try:
            from PyQt5.QtWebEngineWidgets import QWebEngineView
            from PyQt5.QtWidgets import QApplication
            from PyQt5.QtCore import QTimer, QEventLoop, QUrl
            from PyQt5.QtGui import QPixmap
            import os
            
            # Geçici HTML dosyası oluştur
            html_file_path = os.path.join(temp_dir, f"{filename}.html")
            optimized_html = self._optimize_html_for_capture(html_content)
            
            with open(html_file_path, 'w', encoding='utf-8') as f:
                f.write(optimized_html)
            
            # QWebEngineView oluştur
            web_view = QWebEngineView()
            web_view.setFixedSize(1200, 1600)
            web_view.show()
            
            # HTML'i yükle
            web_view.load(QUrl.fromLocalFile(html_file_path))
            
            # Yükleme tamamlanana kadar bekle
            loop = QEventLoop()
            web_view.loadFinished.connect(loop.quit)
            
            timer = QTimer()
            timer.timeout.connect(loop.quit)
            timer.start(8000)
            
            loop.exec_()
            timer.stop()
            
            # Render için bekle
            QApplication.processEvents()
            import time
            time.sleep(1)
            
            # Ekran görüntüsü al
            image_path = os.path.join(temp_dir, f"{filename}.png")
            pixmap = web_view.grab()
            
            if not pixmap.isNull() and pixmap.width() > 100:
                success = pixmap.save(image_path, "PNG", 90)
                web_view.hide()
                web_view.deleteLater()
                
                if success and os.path.exists(image_path):
                    print(f"QWebEngineView ile başarılı: {image_path}")
                    return image_path
            
            web_view.hide()
            web_view.deleteLater()
            return None
                
        except Exception as e:
            print(f"QWebEngineView hatası: {str(e)}")
            return None
    
    def _optimize_html_for_capture(self, html_content):
        """HTML içeriğini görsel yakalama için optimize et"""
        try:
            # CSS optimizasyonları ekle
            optimized_html = html_content.replace(
                '<head>',
                '''<head>
                <style>
                    body { 
                        margin: 20px !important; 
                        padding: 20px !important; 
                        font-family: 'Segoe UI', Arial, sans-serif !important;
                        background: white !important;
                        zoom: 0.8;
                    }
                    .container { 
                        max-width: 1100px !important; 
                        margin: 0 auto !important;
                        box-shadow: none !important;
                    }
                    table { 
                        border-collapse: collapse !important; 
                        width: 100% !important;
                        font-size: 12px !important;
                    }
                    th, td { 
                        border: 1px solid #ddd !important; 
                        padding: 8px !important;
                        text-align: left !important;
                    }
                    th { 
                        background-color: #3498db !important; 
                        color: white !important;
                        font-weight: bold !important;
                    }
                    .metric-card {
                        display: inline-block !important;
                        margin: 10px !important;
                        padding: 15px !important;
                        border: 2px solid #3498db !important;
                        border-radius: 8px !important;
                        text-align: center !important;
                        min-width: 150px !important;
                    }
                    .risk-banner {
                        padding: 15px !important;
                        margin: 10px 0 !important;
                        border-radius: 5px !important;
                        text-align: center !important;
                        font-weight: bold !important;
                        font-size: 16px !important;
                    }
                    h1, h2, h3 {
                        color: #2c3e50 !important;
                        margin: 15px 0 !important;
                    }
                </style>'''
            )
            
            return optimized_html
            
        except Exception as e:
            print(f"HTML optimizasyon hatası: {str(e)}")
            return html_content
    
    def _set_cell_background_color(self, cell, color):
        """Word tablosu hücresine arka plan rengi ekle"""
        try:
            from docx.oxml.shared import OxmlElement, qn
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            # RGB değerlerini hex'e çevir
            hex_color = f"{color.r:02x}{color.g:02x}{color.b:02x}"
            
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except:
            pass  # Renk ayarlanamadıysa devam et
    
    def _convert_html_to_word(self, doc, html_content):
        """HTML içeriğini Word formatına çevir"""
        try:
            from bs4 import BeautifulSoup
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # HTML'i parse et
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # CSS stillerini temizle (sadece içeriği al)
            for style_tag in soup.find_all('style'):
                style_tag.decompose()
            
            # Ana container'ı bul
            container = soup.find('div', class_='container') or soup.find('body') or soup
            
            # İçeriği işle
            self._process_html_element(doc, container)
            
        except Exception as e:
            # HTML parse edilemezse düz metin olarak ekle
            clean_text = re.sub('<[^<]+?>', '', html_content)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            clean_text = clean_text.replace('&nbsp;', ' ')
            clean_text = clean_text.replace('&amp;', '&')
            clean_text = clean_text.replace('&lt;', '<')
            clean_text = clean_text.replace('&gt;', '>')
            
            if clean_text and len(clean_text) > 20:
                doc.add_paragraph(clean_text)
            else:
                doc.add_paragraph("Bu analiz için detay rapor içeriği bulunamadı.")
    
    def _process_html_element(self, doc, element):
        """HTML elementini Word'e çevir"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            if element.name == 'table':
                self._convert_html_table_to_word(doc, element)
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text().strip(), level=min(level, 3))
                if level <= 2:
                    heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    para = doc.add_paragraph(text)
                    # Özel sınıflar için stil uygula
                    if 'risk-banner' in element.get('class', []):
                        para.runs[0].font.bold = True
                        para.runs[0].font.size = Pt(14)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif 'metric-value' in element.get('class', []):
                        para.runs[0].font.bold = True
                        para.runs[0].font.size = Pt(18)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element.name == 'div':
                # Div içindeki elementleri işle
                for child in element.children:
                    if hasattr(child, 'name') and child.name:
                        self._process_html_element(doc, child)
                    elif isinstance(child, str) and child.strip():
                        # Düz metin
                        text = child.strip()
                        if text and len(text) > 3:
                            doc.add_paragraph(text)
            else:
                # Diğer elementler için metin içeriğini al
                text = element.get_text().strip()
                if text and len(text) > 10:
                    doc.add_paragraph(text)
                    
        except Exception as e:
            # Element işlenemezse metin olarak ekle
            text = element.get_text().strip() if hasattr(element, 'get_text') else str(element).strip()
            if text and len(text) > 10:
                doc.add_paragraph(text)
    
    def _convert_html_table_to_word(self, doc, html_table):
        """HTML tablosunu Word tablosuna çevir"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            rows = html_table.find_all('tr')
            if not rows:
                return
            
            # Tablo boyutunu belirle
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            
            # Word tablosu oluştur
            word_table = doc.add_table(rows=len(rows), cols=max_cols)
            word_table.style = 'Table Grid'
            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Satırları doldur
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                word_row = word_table.rows[i]
                
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        word_cell = word_row.cells[j]
                        cell_text = cell.get_text().strip()
                        word_cell.text = cell_text
                        
                        # Hücre formatı
                        paragraph = word_cell.paragraphs[0]
                        
                        # Başlık satırı formatı
                        if cell.name == 'th' or i == 0:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Mavi arka plan
                            self._set_cell_background_color(word_cell, RGBColor(52, 152, 219))
                        else:
                            # Veri satırları
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Sayısal değerler için sağa hizala
                            if cell_text.replace(',', '').replace('.', '').replace('%', '').replace('-', '').isdigit():
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            # Zebra çizgili tablo için alternatif renk
                            if i % 2 == 0:
                                self._set_cell_background_color(word_cell, RGBColor(248, 249, 250))
            
            doc.add_paragraph()  # Tablo sonrası boşluk
            
        except Exception as e:
            # Tablo çevrilemezse basit metin olarak ekle
            table_text = html_table.get_text().strip()
            if table_text:
                doc.add_paragraph(f"📊 Tablo Verisi:\n{table_text}")

    def check_tedarikci_beyan_kontrol(self):
        """Tedarikçi Beyan Kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Tedarikçi beyan kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Tedarikçi beyan kontrol fonksiyonunu çağır
            check_result = check_tedarikci_beyan_kontrol(self.current_df)
            
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            if not check_result['success']:
                QMessageBox.warning(self, "Uyarı", check_result['message'])
                self.status_label.setText("Tedarikçi beyan kontrol analizi tamamlandı")
                self.progress_bar.setVisible(False)
                return
            
            # Sonuçları göster
            if check_result['data'].empty:
                QMessageBox.information(self, "Bilgi", check_result['message'])
                self.status_label.setText("Tedarikçi beyan kontrol analizi tamamlandı")
                self.progress_bar.setVisible(False)
                return
            
            # HTML rapor oluştur
            html_content = self._generate_tedarikci_beyan_html_report(check_result)
            
            # Sonuçları CheckResultsWidget'a ekle
            results = {"Tedarikçi Beyan Kontrol Analizi": check_result}
            self.check_results_widget.add_results(results, html_content)
            
            # Analizler sekmesine geç
            self.tabs.setCurrentIndex(1)
            
            self.progress_bar.setValue(100)
            self.status_label.setText(f"Tedarikçi beyan kontrol analizi tamamlandı: {len(check_result['data'])} kayıt analiz edildi")
            
            # Progress bar'ı gizle
            QTimer.singleShot(2000, lambda: self.progress_bar.setVisible(False))
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Tedarikçi beyan kontrolü sırasında hata: {str(e)}")
            self.status_label.setText("Tedarikçi beyan kontrolünde hata oluştu")
            self.progress_bar.setVisible(False)

    def _generate_tedarikci_beyan_html_report(self, result):
        """Tedarikçi beyan kontrol analizi için HTML rapor oluştur"""
        try:
            df = result['data']
            summary = result.get('summary', {})
            
            html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Tedarikçi Beyan Kontrol Analizi</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                    .risk-high {{ background-color: #ffebee; }}
                    .risk-medium {{ background-color: #fff3e0; }}
                    .risk-low {{ background-color: #e8f5e8; }}
                </style>
            </head>
            <body>
                <h1>Tedarikçi Beyan Kontrol Analizi</h1>
                <p>Toplam {len(df)} kayıt analiz edildi.</p>
                <table>
                    <tr>
                        <th>Beyanname No</th>
                        <th>GTIP</th>
                        <th>Menşe Ülke</th>
                        <th>IGV Ödendi mi?</th>
                        <th>Tedarikçi Beyanı Var mı?</th>
                        <th>Risk Durumu</th>
                    </tr>
            """
            
            for _, row in df.head(100).iterrows():
                risk_class = "risk-high" if row['Risk_durumu'] == 'Yüksek Risk' else "risk-medium" if row['Risk_durumu'] == 'Orta Risk' else "risk-low"
                
                html += f"""
                    <tr class="{risk_class}">
                        <td>{row['Beyanname_no']}</td>
                        <td>{row['Gtip']}</td>
                        <td>{row['Mensei_ulke']}</td>
                        <td>{row['IGV_odendi_mi']}</td>
                        <td>{row['Tedarikci_beyan_var_mi']}</td>
                        <td>{row['Risk_durumu']}</td>
                    </tr>
                """
            
            html += """
                </table>
            </body>
            </html>
            """
            
            return html
            
        except Exception as e:
            return f"<html><body><h1>Rapor Oluşturma Hatası</h1><p>{str(e)}</p></body></html>"

    



    def check_kkdf_kontrol(self):
        """KKDF Kontrol - KKDF yükümlülüğü olan işlemlerde KKDF beyan kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("KKDF kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # KKDF kontrol fonksiyonunu çağır
            kkdf_check = check_kkdf_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if kkdf_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", kkdf_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(kkdf_check["message"])
                return
            
            result = {
                "KKDF Kontrol Analizi": kkdf_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in kkdf_check:
                self.check_results_widget.set_html_report(kkdf_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in kkdf_check and "summary" in kkdf_check:
                self.check_results_widget.show_details(kkdf_check["data"])
                self.check_results_widget.show_summary(kkdf_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(kkdf_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"KKDF kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

    def check_gozetim_kontrol(self):
        """Gözetim Kontrol - Gözetim kapsamındaki GTİP kodları için minimum kıymet eşik kontrolü"""
        if self.current_df is None:
            QMessageBox.warning(self, "Uyarı", "Veri yüklenmedi")
            return
        
        try:
            # İşlem başladığını göster
            self.status_label.setText("Gözetim kontrol analizi çalışıyor...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(10)
            QApplication.processEvents()
            
            # Gözetim kontrol fonksiyonunu çağır
            gozetim_check = check_gozetim_kontrol(self.current_df)
            
            # İşlem tamamlandığını göster
            self.progress_bar.setValue(60)
            QApplication.processEvents()
            
            if gozetim_check["status"] == "error":
                QMessageBox.warning(self, "Uyarı", gozetim_check["message"])
                self.progress_bar.setVisible(False)
                self.status_label.setText(gozetim_check["message"])
                return
            
            result = {
                "Gözetim Kontrol Analizi": gozetim_check
            }
            
            # Sonuçları widget'a aktar
            self.check_results_widget.set_check_results(result)
            
            # İlerleme çubuğunu güncelle
            self.progress_bar.setValue(80)
            QApplication.processEvents()
            
            # HTML raporunu göster
            if "html_report" in gozetim_check:
                self.check_results_widget.set_html_report(gozetim_check["html_report"])
                self.check_results_widget.tabs.setCurrentIndex(3)  # Görsel Rapor sekmesine geçiş
                
            # Tablo ve özet verilerini göster
            if "data" in gozetim_check and "summary" in gozetim_check:
                self.check_results_widget.show_details(gozetim_check["data"])
                self.check_results_widget.show_summary(gozetim_check["summary"])
            
            # Durum bilgisini güncelle
            self.status_label.setText(gozetim_check["message"])
            
            # İşlem tamamlandı
            self.progress_bar.setValue(100)
            QApplication.processEvents()
            self.progress_bar.setVisible(False)
        
        except Exception as e:
            # Hata durumunda kullanıcıyı bilgilendir
            error_msg = f"Gözetim kontrol analizi sırasında hata: {str(e)}"
            print(error_msg)
            QMessageBox.critical(self, "Hata", error_msg)
            self.status_label.setText("Hata oluştu")
            self.progress_bar.setVisible(False)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = CustomsCheckApp()
    window.show()
    sys.exit(app.exec_()) 